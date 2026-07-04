#!/usr/bin/env python3
"""Generate BZPM Executive Presentation Package v2.1 RU — presentation Excel + conclusions Word."""

from __future__ import annotations

import re
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

BASE = Path(__file__).resolve().parent.parent
OUT = Path(__file__).resolve().parent
REGISTRY_MD = BASE / "BZPM-COMPETITOR-REGISTRY-v2.md"
MASTER_MD = BASE / "BZPM-MARKET-INTELLIGENCE-MASTER-REPORT-v1.md"
INSIGHTS_MD = BASE / "BZPM-OPERATOR-INSIGHTS-v1.md"

GEN_DATE = date.today().isoformat()
VERSION = "v2.1 RU"
VERSION_SHORT = "v2.1"

STATUS_RU = {
    "Approved": "Утверждённый реестр",
    "Strong Expansion": "Сильные кандидаты на расширение",
    "Possible Expansion": "Возможные кандидаты",
    "Deferred": "Отложено",
    "Excluded": "Исключено",
}

TIER_RU = {
    "A": "A — Прямые конкуренты / OEM",
    "B": "B — Федеральные лидеры",
    "C": "C — Стратегическое региональное покрытие",
    "D": "D — Агрегаторы",
    "E": "E — Отраслевые маркетплейсы",
    "F": "F — Международные референсы",
}

# Consulting palette
NAVY = "1F4E79"
NAVY_LIGHT = "2E75B6"
ACCENT = "4472C4"
GOLD = "C9A227"
WHITE = "FFFFFF"
LIGHT_BG = "F2F7FB"
CARD_BG = "E8F0FE"
MUTED = "666666"

NAVY_FILL = PatternFill("solid", fgColor=NAVY)
ACCENT_FILL = PatternFill("solid", fgColor=ACCENT)
GOLD_FILL = PatternFill("solid", fgColor=GOLD)
LIGHT_FILL = PatternFill("solid", fgColor=LIGHT_BG)
CARD_FILL = PatternFill("solid", fgColor=CARD_BG)
WHITE_FILL = PatternFill("solid", fgColor=WHITE)

TITLE_FONT = Font(bold=True, size=22, color=NAVY)
SUBTITLE_FONT = Font(bold=True, size=14, color=NAVY)
SECTION_FONT = Font(bold=True, size=16, color=NAVY)
KPI_VALUE_FONT = Font(bold=True, size=28, color=NAVY)
KPI_LABEL_FONT = Font(size=11, color=MUTED)
BODY_FONT = Font(size=11, color="333333")
THIN = Side(style="thin", color="B4B4B4")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def parse_md_table(text: str, start_marker: str, end_marker: str | None = None) -> list[list[str]]:
    idx = text.find(start_marker)
    if idx == -1:
        return []
    chunk = text[idx:]
    if end_marker:
        end = chunk.find(end_marker, len(start_marker))
        if end != -1:
            chunk = chunk[:end]
    rows: list[list[str]] = []
    for line in chunk.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            if rows and not line.startswith("|"):
                break
            continue
        if re.match(r"^\|\s*[-:]+\s*\|", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def normalize_status(raw: str) -> str:
    s = raw.strip()
    if s == "Approved":
        return "Approved"
    if "Strong Expansion" in s:
        return "Strong Expansion"
    if "Possible Expansion" in s:
        return "Possible Expansion"
    if s == "Deferred":
        return "Deferred"
    if s == "Excluded":
        return "Excluded"
    return s


def load_registry() -> list[dict]:
    text = REGISTRY_MD.read_text(encoding="utf-8")
    rows = parse_md_table(text, "## TASK 3 — BZPM Competitor Registry v2", "## TASK 5")
    entities = []
    for row in rows[1:]:
        if len(row) < 9:
            continue
        entities.append(
            {
                "ID": row[0],
                "Company": row[1],
                "Website": row[2],
                "Tier": row[3],
                "Geography": row[4],
                "Coverage Zone": row[5],
                "Type": row[6],
                "Source Waves": row[7],
                "Status": normalize_status(row[8]),
            }
        )
    return entities


def set_col_widths(ws, widths: dict[str, float]):
    for col, w in widths.items():
        ws.column_dimensions[col].width = w


def merge_card(ws, cell_range: str, fill: PatternFill | None = None):
    if fill:
        start, end = cell_range.split(":")
        from openpyxl.utils import range_boundaries
        min_col, min_row, max_col, max_row = range_boundaries(cell_range)
        for r in range(min_row, max_row + 1):
            for c in range(min_col, max_col + 1):
                cell = ws.cell(r, c)
                cell.fill = fill
                cell.border = BORDER
    ws.merge_cells(cell_range)


def add_kpi_card(ws, row: int, col: int, label: str, value, subtitle: str = ""):
    """KPI card block — value + label in styled cells."""
    c1 = get_column_letter(col)
    c2 = get_column_letter(col + 2)
    for r in range(row, row + 4):
        for c in range(col, col + 3):
            cell = ws.cell(r, c)
            cell.fill = CARD_FILL
            cell.border = BORDER
    ws.merge_cells(f"{c1}{row}:{c2}{row}")
    ws[f"{c1}{row}"].value = str(value)
    ws[f"{c1}{row}"].font = KPI_VALUE_FONT
    ws[f"{c1}{row}"].alignment = Alignment(horizontal="center", vertical="center")
    ws.merge_cells(f"{c1}{row + 1}:{c2}{row + 2}")
    ws[f"{c1}{row + 1}"].value = label.replace("\n", " ")
    ws[f"{c1}{row + 1}"].font = KPI_LABEL_FONT
    ws[f"{c1}{row + 1}"].alignment = Alignment(horizontal="center", wrap_text=True, vertical="top")
    if subtitle:
        ws.merge_cells(f"{c1}{row + 3}:{c2}{row + 3}")
        ws[f"{c1}{row + 3}"].value = subtitle
        ws[f"{c1}{row + 3}"].font = Font(size=9, italic=True, color=MUTED)
        ws[f"{c1}{row + 3}"].alignment = Alignment(horizontal="center", wrap_text=True)


def add_bar_chart(ws, title: str, data_col: int, cat_col: int, start: int, end: int, anchor: str):
    chart = BarChart()
    chart.type = "col"
    chart.title = title
    chart.style = 10
    chart.width = 18
    chart.height = 12
    chart.y_axis.title = "Количество"
    data = Reference(ws, min_col=data_col, min_row=start, max_row=end)
    cats = Reference(ws, min_col=cat_col, min_row=start + 1, max_row=end)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    ws.add_chart(chart, anchor)


def add_pie_chart(ws, title: str, data_col: int, cat_col: int, start: int, end: int, anchor: str):
    chart = PieChart()
    chart.title = title
    chart.style = 10
    chart.width = 16
    chart.height = 12
    data = Reference(ws, min_col=data_col, min_row=start, max_row=end)
    cats = Reference(ws, min_col=cat_col, min_row=start + 1, max_row=end)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    ws.add_chart(chart, anchor)


def sheet_header(ws, title: str, subtitle: str = ""):
    ws.merge_cells("A1:H1")
    ws["A1"] = title
    ws["A1"].font = TITLE_FONT
    ws["A1"].alignment = Alignment(vertical="center")
    if subtitle:
        ws.merge_cells("A2:H2")
        ws["A2"] = subtitle
        ws["A2"].font = Font(size=11, italic=True, color=MUTED)
    ws.row_dimensions[1].height = 36


def add_commentary(ws, row: int, text: str, merge_range: str = "A:H"):
    start_col, end_col = merge_range.split(":")
    ws.merge_cells(f"{start_col}{row}:{end_col}{row}")
    cell = ws[f"{start_col}{row}"]
    cell.value = text
    cell.font = Font(size=10, italic=True, color=MUTED)
    cell.alignment = Alignment(wrap_text=True, vertical="top")
    cell.fill = LIGHT_FILL
    for c in range(ord(start_col) - 64, ord(end_col) - 63):
        ws.cell(row, c).border = BORDER
    ws.row_dimensions[row].height = max(36, 14 * (1 + len(text) // 90))


def build_excel(entities: list[dict]) -> list[str]:
    wb = Workbook()
    charts_created: list[str] = []
    approved = [e for e in entities if e["Status"] == "Approved"]
    status_counts = Counter(e["Status"] for e in entities)

    # --- 01 Обложка ---
    ws = wb.active
    ws.title = "01 Обложка"
    set_col_widths(ws, {"A": 4, "B": 18, "C": 18, "D": 18, "E": 18, "F": 4})
    for r in range(3, 16):
        for c in range(2, 6):
            ws.cell(r, c).fill = LIGHT_FILL
            ws.cell(r, c).border = BORDER
    cover_lines = [
        (5, "BZPM Market Intelligence", Font(bold=True, size=28, color=NAVY)),
        (7, "Пакет презентации для руководства", Font(size=16, color=ACCENT)),
        (9, "Исследование рынка нейтрального оборудования для общепита (ЗПМ)", Font(size=12, color=MUTED)),
        (12, f"Версия {VERSION}  ·  {GEN_DATE}", Font(size=11, color=MUTED)),
        (14, "Website Factory  ·  MARS", Font(bold=True, size=11, color=NAVY)),
    ]
    for row_num, text, font in cover_lines:
        ws.merge_cells(f"B{row_num}:E{row_num}")
        cell = ws.cell(row_num, 2)
        cell.value = text
        cell.font = font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # --- 02 Цели исследования ---
    ws2 = wb.create_sheet("02 Цели")
    sheet_header(ws2, "Цели исследования", "Зачем проводилось исследование рынка ЗПМ")
    add_commentary(
        ws2,
        3,
        "Исследование выполнено перед разработкой нового сайта ЗПМ (Барнаул). "
        "Его задача — зафиксировать конкурентное поле, географию спроса и практики отрасли, "
        "чтобы проектирование сайта опиралось на проверенные данные, а не на предположения.",
    )
    goals = [
        ("Контекст", "ЗПМ (Барнаул) — производитель нейтрального оборудования для общепита. "
         "Перед разработкой нового сайта необходимо понять конкурентное поле, географию рынка и лучшие практики отрасли."),
        ("Задача 1", "Сформировать полную карту конкурентов в сегменте нейтрального оборудования — от OEM-производителей до федеральных дистрибьюторов и агрегаторов."),
        ("Задача 2", "Определить приоритетных конкурентов для глубокого анализа (Tier A–F) с учётом географии Сибири, Урала и Дальнего Востока."),
        ("Задача 3", "Выявить региональные пробелы покрытия — Барнаул, Сибирь, Урал, Дальний Восток, Казахстан, Беларусь."),
        ("Задача 4", "Зафиксировать видимость конкурентов в поисковой выдаче (SERP) по ключевым запросам."),
        ("Задача 5", "Собрать операторские наблюдения по UX каталога, фильтрации, карточкам товара и коммерческим блокам."),
        ("Вопросы", "Кто конкуренты ЗПМ? · Какие типы игроков на рынке? · Кто доминирует в поиске? · "
         "Какие UX-паттерны используют лидеры? · Что применить при разработке нового сайта?"),
    ]
    r = 5
    for title, text in goals:
        ws2.merge_cells(f"A{r}:H{r}")
        ws2[f"A{r}"] = title
        ws2[f"A{r}"].font = SECTION_FONT
        r += 1
        ws2.merge_cells(f"A{r}:H{r}")
        ws2[f"A{r}"] = text
        ws2[f"A{r}"].font = BODY_FONT
        ws2[f"A{r}"].alignment = Alignment(wrap_text=True)
        ws2.row_dimensions[r].height = 48 if len(text) > 80 else 28
        r += 2
    set_col_widths(ws2, {"A": 14, "B": 14, "C": 14, "D": 14, "E": 14, "F": 14, "G": 14, "H": 14})

    # --- 03 Методология ---
    ws3 = wb.create_sheet("03 Методология")
    sheet_header(ws3, "Методология исследования", "Последовательность волн W1 → W3Y")
    add_commentary(
        ws3,
        3,
        "Исследование проводилось поэтапно — от общей карты рынка до фиксации UX-наблюдений. "
        "Коды W1–W3Y — внутренние обозначения этапов; ниже каждая волна расшифрована простым языком.",
    )
    waves = [
        ("W1", "Картирование рынка", "Определение позиционирования ЗПМ и базового пути покупателя", "Базовая карта сегмента и точек контакта"),
        ("W2", "Поиск конкурентов", "Сбор кандидатов по тирам A–F", "80 кандидатов в discovery universe"),
        ("W2.5", "Приоритизация", "Отбор и отсев кандидатов", "Core Research Set — 46 Approved; 34 отложено"),
        ("W3", "Реестр конкурентов", "Формирование реестра с COMP-BZPM ID", "46 сущностей; оценка покрытия; 13 review flags"),
        ("W3R", "Региональное усиление", "Усиление по Сибири, Уралу, ДВ, KZ, BY", "38 региональных строк; пулы Strong/Possible/Low"),
        ("W3S", "Поисковая видимость (SERP)", "Анализ присутствия в поисковой выдаче", "Матрица 20 запросов; 40 доменов; 27 новых кандидатов"),
        ("W3X", "Консолидация реестра", "Сведение реестра v2 + Master Report", "126 канонических сущностей после dedup"),
        ("W3Y", "Операторские наблюдения UX", "Фиксация наблюдений при ручном просмотре сайтов", "7 highlights; 6 benchmark; 7 FIM markers"),
    ]
    ws3["A4"] = "Этап"
    ws3["B4"] = "Название"
    ws3["C4"] = "Что делалось"
    ws3["D4"] = "Результат"
    for c in "ABCD":
        ws3[f"{c}4"].font = Font(bold=True, color=WHITE)
        ws3[f"{c}4"].fill = NAVY_FILL
    r = 5
    for w, name, what, result in waves:
        ws3[f"A{r}"] = w
        ws3[f"B{r}"] = name
        ws3[f"C{r}"] = what
        ws3[f"D{r}"] = result
        for c in "ABCD":
            ws3[f"{c}{r}"].font = BODY_FONT
            ws3[f"{c}{r}"].alignment = Alignment(wrap_text=True, vertical="top")
            ws3[f"{c}{r}"].fill = LIGHT_FILL if r % 2 == 0 else WHITE_FILL
            ws3[f"{c}{r}"].border = BORDER
        ws3.row_dimensions[r].height = 36
        r += 1
    # Pipeline arrows column
    ws3["F5"] = "W1"
    for i, (w, _, _, _) in enumerate(waves[1:], start=6):
        ws3[f"F{i}"] = "↓"
        ws3[f"G{i}"] = w
        ws3[f"F{i}"].alignment = Alignment(horizontal="center")
        ws3[f"G{i}"].font = Font(bold=True, color=ACCENT)
    set_col_widths(ws3, {"A": 8, "B": 22, "C": 38, "D": 38, "F": 4, "G": 8})

    # --- 04 Масштаб ---
    ws4 = wb.create_sheet("04 Масштаб")
    sheet_header(ws4, "Общий масштаб исследования", "Ключевые показатели по утверждённым данным")
    add_commentary(
        ws4,
        3,
        "126 сущностей (entities) — вся найденная и нормализованная конкурентная вселенная после консолидации. "
        "46 Approved — проверенное ядро реестра, на которое опираются выводы и приоритеты для сайта ЗПМ.",
    )
    kpis = [
        ("126", "Канонических\nсущностей", "после dedup W3X"),
        ("46", "Утверждённый\nреестр (Approved)", "базовый слой W3"),
        ("21", "Сильные кандидаты\nна расширение", "Strong Expansion"),
        ("22", "Возможные\nкандидаты", "Possible Expansion"),
        ("80", "Пул W2\nDiscovery", "сырой пул кандидатов"),
        ("40", "Доменов\nв SERP", "отслеживалось W3S"),
        ("27", "Новых доменов\nиз SERP", "не были в реестре"),
        ("20", "Поисковых\nзапросов", "матрица W3S"),
    ]
    positions = [(4, 1), (4, 4), (4, 7), (10, 1), (10, 4), (10, 7), (16, 1), (16, 4)]
    for (val, label, sub), (row, col) in zip(kpis, positions):
        add_kpi_card(ws4, row, col, label, val, sub)
    set_col_widths(ws4, {"A": 12, "B": 12, "C": 12, "D": 12, "E": 12, "F": 12, "G": 12, "H": 4})

    # --- 05 География ---
    ws5 = wb.create_sheet("05 География")
    sheet_header(ws5, "География исследования", "Регионы, страны и зоны покрытия")
    add_commentary(
        ws5,
        3,
        "Для ЗПМ критичны Сибирь, Урал и Дальний Восток — это первичная география продаж и штаб-квартира (Барнаул). "
        "Исследование фиксирует, где конкуренты представлены сильно, частично или с пробелами в покрытии.",
    )
    geo_counter: Counter = Counter()
    for e in entities:
        g = e["Geography"].split("(")[0].strip()
        if g.startswith("Russia"):
            geo_counter["Россия"] += 1
        elif "Kazakhstan" in g:
            geo_counter["Казахстан"] += 1
        elif "Belarus" in g:
            geo_counter["Беларусь"] += 1
        elif "International" in g:
            geo_counter["Международные"] += 1
    ws5.append([])
    ws5.append(["География HQ", "Сущностей"])
    hdr = ws5.max_row
    for c in range(1, 3):
        ws5.cell(hdr, c).font = Font(bold=True, color=WHITE)
        ws5.cell(hdr, c).fill = NAVY_FILL
    for g, cnt in sorted(geo_counter.items(), key=lambda x: -x[1]):
        ws5.append([g, cnt])
    add_bar_chart(ws5, "География HQ (все сущности)", 2, 1, hdr, ws5.max_row, "D4")
    charts_created.append("География HQ — bar chart (05 География)")

    zones = [
        ("Сибирь", 14, "+6 strong expansion"),
        ("Урал", 13, "+2 strong expansion"),
        ("Дальний Восток", 8, "+3 strong expansion (косвенно)"),
        ("Европейская Россия", 28, "+5 federal/OEM"),
        ("Казахстан", 10, "+4 strong expansion"),
        ("Беларусь", 10, "+2 strong expansion"),
        ("СНГ", 12, "покрытие СНГ"),
        ("Международные", 10, "референсы Tier F"),
    ]
    zr = ws5.max_row + 3
    ws5.cell(zr, 1, "Зона покрытия").font = SECTION_FONT
    zr += 1
    ws5.append(["Зона покрытия", "Approved (W3)", "Добавления Strong expansion"])
    hdr2 = ws5.max_row
    for c in range(1, 4):
        ws5.cell(hdr2, c).font = Font(bold=True, color=WHITE)
        ws5.cell(hdr2, c).fill = ACCENT_FILL
    for z, a, s in zones:
        ws5.append([z, a, s])
    set_col_widths(ws5, {"A": 22, "B": 14, "C": 28, "D": 4, "E": 4, "F": 4})

    # --- 06 Классификация ---
    ws6 = wb.create_sheet("06 Классификация")
    sheet_header(ws6, "Классификация компаний", "Tier · статус · тип компании")
    add_commentary(
        ws6,
        3,
        "Tier A–F — шкала роли на рынке: от прямых OEM-конкурентов (A) до международных референсов (F). "
        "Статусы реестра показывают, какие компании в утверждённом ядре, а какие — кандидаты на расширение или отложены.",
    )
    tier_labels = TIER_RU
    tier_counter = Counter(e["Tier"] for e in approved if e["Tier"] in tier_labels)
    ws6.append(["Tier", "Описание", "Approved"])
    hdr = ws6.max_row
    for c in range(1, 4):
        ws6.cell(hdr, c).font = Font(bold=True, color=WHITE)
        ws6.cell(hdr, c).fill = NAVY_FILL
    for t in "ABCDEF":
        ws6.append([t, tier_labels[t], tier_counter.get(t, 0)])
    add_pie_chart(ws6, "Классификация по Tier (Approved)", 3, 2, hdr, ws6.max_row, "E4")
    charts_created.append("Классификация по Tier — pie (06 Классификация)")

    sr = ws6.max_row + 3
    ws6.cell(sr, 1, "Статус в реестре").font = SECTION_FONT
    sr += 1
    ws6.append(["Статус", "Количество"])
    hdr3 = ws6.max_row
    for c in range(1, 3):
        ws6.cell(hdr3, c).font = Font(bold=True, color=WHITE)
        ws6.cell(hdr3, c).fill = ACCENT_FILL
    for st in ("Approved", "Strong Expansion", "Possible Expansion", "Deferred", "Excluded"):
        ws6.append([STATUS_RU.get(st, st), status_counts.get(st, 0)])
    add_pie_chart(ws6, "Статусы реестра (все 126)", 2, 1, hdr3, ws6.max_row, "E18")
    charts_created.append("Статусы реестра — pie (06 Классификация)")
    set_col_widths(ws6, {"A": 12, "B": 38, "C": 12, "D": 4, "E": 4})

    # --- 07 Распределение конкурентов ---
    ws7 = wb.create_sheet("07 Распределение")
    sheet_header(ws7, "Распределение конкурентов", "По типам · регионам · статусам")
    add_commentary(
        ws7,
        3,
        "Диаграммы показывают структуру исследованного рынка — не объём продаж и не финансовый рейтинг. "
        "Сравнивается состав конкурентного поля: типы компаний и география утверждённого реестра.",
    )
    type_counter = Counter(e["Type"] for e in approved)
    ws7.append(["Тип компании (Approved)", "Количество"])
    hdr = ws7.max_row
    for c in range(1, 3):
        ws7.cell(hdr, c).font = Font(bold=True, color=WHITE)
        ws7.cell(hdr, c).fill = NAVY_FILL
    for t, cnt in sorted(type_counter.items(), key=lambda x: -x[1]):
        ws7.append([t, cnt])
    add_bar_chart(ws7, "Типы компаний (Approved)", 2, 1, hdr, ws7.max_row, "D4")
    charts_created.append("Типы компаний — bar (07 Распределение)")

    gr = ws7.max_row + 3
    ws7.cell(gr, 1, "География утверждённого реестра").font = SECTION_FONT
    gr += 1
    geo_app = Counter()
    for e in approved:
        g = e["Geography"]
        if g.startswith("Russia"):
            geo_app["Россия"] += 1
        elif "Kazakhstan" in g:
            geo_app["Казахстан"] += 1
        elif "Belarus" in g:
            geo_app["Беларусь"] += 1
        elif "International" in g:
            geo_app["Международные"] += 1
    ws7.append(["География", "Approved"])
    hdr2 = ws7.max_row
    for c in range(1, 3):
        ws7.cell(hdr2, c).font = Font(bold=True, color=WHITE)
        ws7.cell(hdr2, c).fill = ACCENT_FILL
    for g, cnt in sorted(geo_app.items(), key=lambda x: -x[1]):
        ws7.append([g, cnt])
    add_bar_chart(ws7, "География Approved", 2, 1, hdr2, ws7.max_row, "D18")
    charts_created.append("География Approved — bar (07 Распределение)")
    set_col_widths(ws7, {"A": 28, "B": 12, "C": 4, "D": 4})

    # --- 08 Benchmark ---
    ws8 = wb.create_sheet("08 Benchmark")
    sheet_header(ws8, "Группа Benchmark", "Список для анализа решений — не рейтинг компаний")
    add_commentary(
        ws8,
        3,
        "Benchmark group — компании, к которым оператор неоднократно возвращался при ручном просмотре. "
        "Это не рейтинг качества и не рекомендация «лучший сайт» — список для изучения конкретных UX-решений.",
    )
    benchmark = [
        ("УЗНМ", "https://zavod-uznm.ru/", "COMP-BZPM-007",
         "OEM peer Tier A; отмечен упрощённый паттерн фильтрации", "Страницы категорий; системы фильтров"),
        ("КЛЕН", "https://www.klenmarket.ru/", "COMP-BZPM-012",
         "Федеральный лидер; паттерн переключателя вида; SERP 7×", "Макеты каталога; страницы категорий"),
        ("ГК Юниторг", "https://www.unitorg.ru/", "CAN-EXP-005",
         "Уральский регионал; двухколоночные карточки в listing", "Страницы категорий; листинг товаров"),
        ("Trapeza", "https://www.trapeza.ru/", "COMP-BZPM-011",
         "Федеральный якорь; плотность информации; SERP 8×", "Каталог; категории; карточки товара"),
        ("Kobor", "https://kobor.ru/", "CAN-EXP-003",
         "Повторное операторское внимание; мультирегиональные подсайты", "Обзор регионального покрытия"),
        ("Комплекс Трейд", "https://kompleks-trade.ru/", "CAN-EXP-004",
         "Повторное операторское внимание; Урал/Сибирь multi-city", "Обзор регионального покрытия"),
    ]
    ws8.append(["Компания", "Сайт", "ID реестра", "Почему в benchmark", "Поверхности для review"])
    hdr = ws8.max_row
    for c in range(1, 6):
        ws8.cell(hdr, c).font = Font(bold=True, color=WHITE)
        ws8.cell(hdr, c).fill = NAVY_FILL
    for row in benchmark:
        ws8.append(list(row))
        for c in range(1, 6):
            cell = ws8.cell(ws8.max_row, c)
            cell.font = BODY_FONT
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = BORDER
    note_row = ws8.max_row + 2
    ws8.merge_cells(f"A{note_row}:E{note_row}")
    note_cell = ws8.cell(note_row, 1)
    note_cell.value = (
        "Критерий отбора: компании, к которым оператор неоднократно возвращался при ручном review (W3Y). "
        "Не формальный рейтинг — attention list для последующих волн анализа (W4+)."
    )
    note_cell.font = Font(italic=True, size=10, color=MUTED)
    note_cell.alignment = Alignment(wrap_text=True)
    set_col_widths(ws8, {"A": 18, "B": 28, "C": 16, "D": 36, "E": 28})

    # --- 09 SERP ---
    ws9 = wb.create_sheet("09 SERP")
    sheet_header(ws9, "Поисковая видимость (SERP)", "W3S — присутствие в поисковой выдаче")
    add_commentary(
        ws9,
        3,
        "SERP-показатели отражают видимость сайта в поисковой выдаче по ключевым запросам сегмента — "
        "не качество компании и не объём продаж. Высокая видимость федеральных игроков не означает слабость регионалов как производителей.",
    )
    serp = [
        ("Trapeza", "https://www.trapeza.ru/", 8, "Approved"),
        ("КЛЕН", "https://www.klenmarket.ru/", 7, "Approved"),
        ("РЕФРО", "https://www.refro.ru/", 5, "Approved"),
        ("Практика (Pectopah)", "https://www.pectopah.ru/", 5, "Strong Expansion"),
        ("Завод Проммаш", "https://prommash.com/", 5, "Strong Expansion"),
        ("Kobor", "https://kobor.ru/", 5, "Strong Expansion"),
        ("Abat", "https://abat.ru/", 4, "Approved"),
        ("Restoll", "https://restoll.ru/", 4, "Approved"),
        ("Finist", "https://f-inox.ru/", 4, "Approved"),
        ("Ресторан Комплект", "https://r-komplekt.ru/", 4, "Strong Expansion"),
        ("КАМИК", "https://kamik-group.ru/", 4, "Strong Expansion"),
        ("Завод МетаКон", "https://zavod-metakon.ru/", 4, "Strong Expansion"),
        ("НОТИС", "https://www.notis.ru/", 4, "Strong Expansion"),
    ]
    ws9.append(["Компания", "Сайт", "Появления в выдаче", "Статус в реестре"])
    hdr = ws9.max_row
    for c in range(1, 5):
        ws9.cell(hdr, c).font = Font(bold=True, color=WHITE)
        ws9.cell(hdr, c).fill = NAVY_FILL
    for row in serp:
        company, site, appearances, status = row
        ws9.append([company, site, appearances, STATUS_RU.get(status, status)])
    add_bar_chart(ws9, "Лидеры поисковой видимости (W3S)", 3, 1, hdr, ws9.max_row, "F4")
    charts_created.append("Лидеры SERP — bar (09 SERP)")
    findings_row = ws9.max_row + 3
    findings = [
        "Исследовалась матрица из 20 ключевых запросов сегмента нейтрального оборудования.",
        "Проанализировано 40 доменов; 27 новых кандидатов не входили в утверждённый реестр (Approved).",
        "Закономерность: федеральное доминирование (Trapeza 8×, КЛЕН 7×) при региональных blind spots.",
        "Региональные игроки (Kobor, НОТИС, Практика) видимы в SERP, но не все в Approved registry.",
        "Примечание: счётчики W3S — curated synthesis; автоматический Yandex parse blocked (SAFE UNKNOWN).",
    ]
    ws9.cell(findings_row, 1, "Закономерности").font = SECTION_FONT
    for i, f in enumerate(findings, start=findings_row + 1):
        ws9.merge_cells(f"A{i}:E{i}")
        ws9[f"A{i}"] = f"• {f}"
        ws9[f"A{i}"].font = BODY_FONT
        ws9[f"A{i}"].alignment = Alignment(wrap_text=True)
    set_col_widths(ws9, {"A": 22, "B": 30, "C": 12, "D": 18, "E": 4, "F": 4})

    # --- 10 Что изучено ---
    ws10 = wb.create_sheet("10 Изучено")
    sheet_header(ws10, "Что было изучено", "Поверхности и объекты анализа")
    add_commentary(
        ws10,
        3,
        "При ручном просмотре сайтов конкурентов фиксировались элементы, влияющие на проектирование каталога и карточек ЗПМ: "
        "навигация, фильтры, плотность информации на листинге, коммерческие блоки и региональная архитектура.",
    )
    surfaces = [
        ("Каталог", "Структура категорий, корневые разделы, навигация по ассортименту нейтрального оборудования"),
        ("Карточки товара (PDP)", "Плотность информации, атрибуты, коммерческие блоки на странице продукта"),
        ("UX / раскладка", "Раскладка listing, двухколоночные карточки, переключатели вида, плотность информации"),
        ("Навигация", "Структура меню, хлебные крошки, переходы каталог → категория → PDP"),
        ("Фильтрация", "UX фильтров — упрощённые vs тяжёлые панели (УЗНМ vs федеральные peers)"),
        ("Коммерческие блоки", "Паттерны цен/CTA, B2B purchase flow, legacy vs modern architecture"),
        ("Контент", "Описания, спецификации, поля для решения на listing без открытия PDP"),
        ("Формы", "Запросы, обратная связь, коммерческие CTA на catalog/PDP surfaces"),
        ("Структура сайта", "Мультигородские подсайты, региональное покрытие, федеральная vs региональная архитектура"),
    ]
    ws10.append(["Область", "Что анализировалось"])
    hdr = ws10.max_row
    for c in range(1, 3):
        ws10.cell(hdr, c).font = Font(bold=True, color=WHITE)
        ws10.cell(hdr, c).fill = NAVY_FILL
    for area, desc in surfaces:
        ws10.append([area, desc])
        for c in range(1, 3):
            cell = ws10.cell(ws10.max_row, c)
            cell.font = BODY_FONT
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = BORDER
    set_col_widths(ws10, {"A": 20, "B": 70})

    # --- 11 Результат ---
    ws11 = wb.create_sheet("11 Результат")
    sheet_header(ws11, "Результат исследования", "Только зафиксированные факты — без финальных рекомендаций")
    add_commentary(
        ws11,
        3,
        "Этот лист — сводка того, что собрано и изучено. Аналитические выводы и рекомендации для сайта — "
        "в документе BZPM Research Conclusions.docx; здесь только factual result без интерпретации.",
    )
    facts = [
        ("Собрано", [
            "126 канонических сущностей после deduplication (W3X)",
            "46 approved registry entities с COMP-BZPM-001…046",
            "21 Strong + 22 Possible expansion candidates",
            "26 deferred + 11 excluded entities",
            "SERP visibility map: 40 domains, top-15 leaders зафиксированы",
            "7 operator highlights + 10 observed patterns (W3Y)",
            "6 Native Benchmark Group companies + 7 FIM markers",
        ]),
        ("Изучено", [
            "Tier structure A–F: OEM, federal, regional, aggregators, marketplaces, international",
            "География: Россия (33 approved), KZ (1), BY (2), International (10)",
            "Региональные зоны: Siberia 14, Ural 13, Far East 8 indirect, Barnaul gap confirmed",
            "UX surfaces: catalog, PDP, filters, view switchers, commercial blocks (operator review)",
            "13 Review Required tier flags — tier ambiguity, not blocking",
        ]),
        ("Построено", [
            "BZPM Competitor Registry v2 (authority markdown)",
            "BZPM Market Intelligence Master Report v1",
            "BZPM Operator Insights v1 (W3Y capture)",
            "Presentation Pack (6 Excel workbooks)",
            "Executive Presentation Package v2.1 RU (этот документ)",
            "Recommended manual review order: 48 items",
        ]),
    ]
    r = 5
    for section, items in facts:
        ws11.merge_cells(f"A{r}:H{r}")
        ws11[f"A{r}"] = section
        ws11[f"A{r}"].font = SECTION_FONT
        r += 1
        for item in items:
            ws11.merge_cells(f"A{r}:H{r}")
            ws11[f"A{r}"] = f"• {item}"
            ws11[f"A{r}"].font = BODY_FONT
            ws11[f"A{r}"].alignment = Alignment(wrap_text=True)
            r += 1
        r += 1
    set_col_widths(ws11, {"A": 14, "B": 14, "C": 14, "D": 14, "E": 14, "F": 14, "G": 14, "H": 14})

    # --- 12 Приложение ---
    ws12 = wb.create_sheet("12 Приложение")
    sheet_header(ws12, "Приложение", "Authority-документы и ссылки на полный реестр")
    add_commentary(
        ws12,
        3,
        "Полный реестр из 126 сущностей и детальная аналитика — в authority markdown родительской папки. "
        "Этот пакет — презентационный слой; для верификации любого утверждения используйте документы ниже.",
    )
    ws12.append(["Документ authority", "Путь", "Роль"])
    hdr = ws12.max_row
    for c in range(1, 4):
        ws12.cell(hdr, c).font = Font(bold=True, color=WHITE)
        ws12.cell(hdr, c).fill = NAVY_FILL
    authorities = [
        ("BZPM-COMPETITOR-REGISTRY-v2.md", "../BZPM-COMPETITOR-REGISTRY-v2.md", "Канонический реестр — 126 сущностей"),
        ("BZPM-MARKET-INTELLIGENCE-MASTER-REPORT-v1.md", "../BZPM-MARKET-INTELLIGENCE-MASTER-REPORT-v1.md", "Базовый master report"),
        ("BZPM-OPERATOR-INSIGHTS-v1.md", "../BZPM-OPERATOR-INSIGHTS-v1.md", "Операторские наблюдения W3Y"),
        ("Presentation Pack", "../presentation-pack/", "6 Excel workbooks — операционный слой"),
        ("Executive Report", "./", "Презентационный Excel + Word с выводами"),
    ]
    for row in authorities:
        ws12.append(list(row))
    ws12.append([])
    ws12.append(["Tier A OEM (Approved)", "Сайт", "ID"])
    hdr2 = ws12.max_row
    for c in range(1, 4):
        ws12.cell(hdr2, c).font = Font(bold=True, color=WHITE)
        ws12.cell(hdr2, c).fill = ACCENT_FILL
    tier_a = [e for e in approved if e["Tier"] == "A"]
    for e in tier_a:
        ws12.append([e["Company"], e["Website"], e["ID"]])
    set_col_widths(ws12, {"A": 36, "B": 42, "C": 36})

    out_path = OUT / "BZPM Market Research.xlsx"
    wb.save(out_path)
    return charts_created


def build_word() -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def heading(text: str, level: int = 1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    def para(text: str, bold: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        return p

    def bullet(text: str):
        doc.add_paragraph(text, style="List Bullet")

    def commentary(text: str):
        p = doc.add_paragraph()
        run = p.add_run("Пояснение: " + text)
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.size = Pt(10)

    # Title page
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BZPM Research Conclusions\n")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    r2 = t.add_run("Аналитические выводы исследования рынка ЗПМ\n")
    r2.font.size = Pt(14)
    r3 = t.add_run(f"Пакет презентации для руководства {VERSION}  ·  {GEN_DATE}\nWebsite Factory · MARS")
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # 1. Введение
    heading("1. Введение")
    para(
        "Настоящий документ содержит аналитические выводы программы BZPM Market Intelligence — "
        "исследования конкурентного поля в сегменте нейтрального оборудования для общепита. "
        "Исследование проводилось в рамках подготовки к разработке нового сайта ЗПМ (Барнаул) "
        "и охватывает волны W1–W3Y: от картирования рынка до фиксации операторских UX-наблюдений."
    )
    para(
        "Документ не дублирует реестр конкурентов и не является рабочей таблицей. "
        "Все выводы основаны на утверждённых authority-материалах: Registry v2, Master Report v1, Operator Insights v1."
    )

    # 2. Картина рынка
    heading("2. Картина рынка")
    para(
        "Рынок нейтрального оборудования для общепита структурирован по шести тирам конкурентов (Tier A–F — "
        "шкала роли на рынке). Ядро — 10 OEM-производителей (Tier A), федеральные лидеры (Tier B, 8 компаний) "
        "и стратегический региональный слой (Tier C, 8 компаний). Дополнительно зафиксированы "
        "агрегаторы (Tier D), отраслевые маркетплейсы (Tier E) и международные референсы (Tier F)."
    )
    bullet("Первичная география ЗПМ: Сибирь, Урал, Дальний Восток (штаб-квартира — Барнаул).")
    bullet("Федеральные поставщики доминируют в поисковой видимости (SERP — присутствие в поисковой выдаче): Trapeza (8 появлений), КЛЕН (7).")
    bullet("Региональные пробелы подтверждены: нет dedicated approved entry для Барнаула; Дальний Восток — PARTIAL.")
    bullet("Каноническая вселенная после консолидации: 126 сущностей; 43 кандидата в expansion pipeline (очередь на расширение реестра).")
    commentary(
        "Для сайта ЗПМ это означает: конкуренция идёт одновременно с производителями (Tier A), "
        "федеральными дистрибьюторами с сильным SEO (Tier B) и региональными игроками. "
        "География Сибири/Урала/ДВ — не второстепенный фактор, а зона, где нужно осознанное позиционирование."
    )

    # 3. Типы производителей
    heading("3. Основные типы производителей")
    para("В утверждённом реестре (Approved registry — 46 проверенных компаний) зафиксированы восемь типов:")
    bullet("OEM Manufacturer (10) — прямые производители нейтрального оборудования, Tier A.")
    bullet("Federal Supplier (5) — федеральные дистрибьюторы с широким ассортиментом и логистикой.")
    bullet("Regional Supplier (8) — региональные игроки с фокусом на Сибирь, Урал, KZ, BY.")
    bullet("Distributor (3), Aggregator (2), Industry Directory (3), Marketplace (5), International OEM (10).")
    para(
        "Для ЗПМ наиболее релевантны Tier A OEM (прямая конкуренция по продукту) и Tier B federal "
        "(конкуренция за внимание покупателя и SERP-позиции). Tier F служит референсом UX и продуктовых стандартов."
    )

    # 4. Ключевые конкуренты
    heading("4. Ключевые конкуренты")
    para("По совокупности tier, SERP-видимости и операторского внимания выделяются:")
    bullet("Abat, Restoinox, Finist, Kroner, BSV-inox, Техно-ТТ, УЗНМ — Tier A OEM peers.")
    bullet("Trapeza — федеральный якорь, максимальная SERP-видимость (8×), высокая плотность информации на листинге.")
    bullet("КЛЕН — федеральный лидер, паттерн переключателя вида каталога, SERP 7×.")
    bullet("Kobor, Комплекс Трейд — Strong Expansion, многократное операторское внимание, мультирегиональные подсайты.")
    bullet("УЗНМ — Tier A OEM с покрытием Урал/Сибирь; упрощённый паттерн фильтрации как дифференциатор.")

    # 5. Benchmark
    heading("5. Benchmark")
    para(
        "Native Benchmark Group (W3Y) — не формальный рейтинг, а benchmark group (список компаний для анализа конкретных UX-решений), "
        "к которым оператор неоднократно возвращался при ручном review. "
        "Шесть компаний: УЗНМ, КЛЕН, ГК Юниторг, Trapeza, Kobor, Комплекс Трейд."
    )
    bullet("УЗНМ — эталон упрощённой фильтрации каталога для OEM peer.")
    bullet("КЛЕН — эталон переключателей вида каталога (grid/list toggle).")
    bullet("ГК Юниторг — эталон двухколоночных карточек на listing.")
    bullet("Trapeza — эталон плотности информации на catalog и PDP surfaces.")
    bullet("Kobor и Комплекс Трейд — эталоны региональной архитектуры (мультигородские подсайты).")
    commentary(
        "Benchmark не отвечает на вопрос «кто лучший на рынке». Он указывает, у каких сайтов "
        "стоит изучить конкретные решения — фильтры, карточки, региональную структуру — при проектировании ЗПМ."
    )

    # 6. Закономерности рынка
    heading("6. Основные закономерности рынка")
    bullet("Федеральное доминирование в SERP при слабой видимости локальных игроков Барнаула.")
    bullet("Региональные OEM/supplier (Kobor, НОТИС, Практика) видимы в поиске, но не все в Approved registry — сигнал для expansion triage.")
    bullet("Архитектура мультигородских подсайтов — стандарт для сильных региональных игроков (Kobor, Комплекс Трейд, Lerius, ФудПром).")
    bullet("Tier C (Strategic Regional Coverage) — критический слой для Сибири/Урала/ДВ, но покрытие неравномерное.")
    bullet("13 Review Required flags в approved registry — неоднозначность tier, не блокирует консолидацию.")

    # 7. UX-находки
    heading("7. UX-находки")
    para("Операторские наблюдения W3Y (без scoring):")
    bullet("Упрощённый паттерн фильтрации (УЗНМ) — меньше групп фильтров, чище presentation vs тяжёлые федеральные панели.")
    bullet("Переключатели вида каталога (КЛЕН) — выбор режима отображения на category/catalog surfaces.")
    bullet("Двухколоночные карточки каталога (ГК Юниторг) — двухколоночная структура внутри product card.")
    bullet("Плотность информации (Trapeza) — больше полей для решения на listing без открытия PDP.")
    bullet("Устаревшая коммерческая архитектура (Энтеро) — legacy commercial blocks и B2B flow patterns.")
    bullet("Навигационные паттерны — формально не тегированы в W3Y beyond listed entities.")
    commentary(
        "Эти находки напрямую влияют на проектирование сайта ЗПМ: что показывать на листинге, "
        "как устроить фильтры OEM-каталога, нужны ли переключатели вида и какая плотность информации ожидается B2B-покупателем."
    )

    # 8. Лучшие практики
    heading("8. Что используют лучшие компании")
    bullet("Высокая плотность информации на listing (Trapeza) — brand, model, section attributes inline.")
    bullet("Переключатели вида для каталога (КЛЕН) — пользовательский выбор layout mode.")
    bullet("Упрощённая фильтрация для OEM-каталогов (УЗНМ) — reduced filter UI.")
    bullet("Двухколоночная структура карточек на listing (ГК Юниторг).")
    bullet("Мультирегиональная архитектура подсайтов (Kobor, Комплекс Трейд, ФудПром).")
    bullet("Плотная структура корня каталога с навигационной и продуктовой информацией per viewport (Trapeza).")

    # 9. Ошибки
    heading("9. Какие ошибки встречаются чаще всего")
    bullet("Устаревшая коммерческая архитектура — legacy pricing/CTA patterns (наблюдение: Энтеро).")
    bullet("«Тонкие» карточки на listing — мало атрибутов, необходимость открывать PDP для базовых решений (контраст с Trapeza).")
    bullet("Тяжёлые панели фильтров без упрощения — федеральные peers vs упрощённый паттерн УЗНМ.")
    bullet("Региональные blind spots — отсутствие dedicated local presence при federal SERP dominance.")
    bullet("Неоднозначность tier — 13 entities с Review Required flags, неоднозначная классификация роли на рынке.")

    # 10. Что использовать
    heading("10. Что необходимо использовать в новом сайте ЗПМ")
    bullet("Плотность информации на catalog/PDP — decision-relevant fields visible per screen (референс: Trapeza).")
    bullet("Упрощённый UX фильтров для OEM-каталога — streamlined filter groups (референс: УЗНМ).")
    bullet("Переключатели вида каталога — grid/list toggle где уместно (референс: КЛЕН).")
    bullet("Структурированные product cards с inline attributes на listing (референс: Trapeza, Юниторг dual-column).")
    bullet("Региональная архитектура покрытия — учёт Сибири, Урала, ДВ в структуре сайта (референс: Kobor, Комплекс Трейд).")
    bullet("Benchmark-driven deep review routing — FIM (Future Investigation Markers, маркеры для углублённого изучения) W3Y-001…007 как entry points для W4+ waves.")
    commentary(
        "Каждый пункт связан с конкретными страницами будущего сайта: каталог, категории, карточки товара, "
        "региональные разделы. При проектировании макетов и прототипов эти паттерны — проверенная отправная точка, "
        "а не абстрактные пожелания."
    )

    # 11. Что нельзя
    heading("11. Что использовать нельзя")
    bullet("Устаревшие паттерны коммерческой архитектуры — legacy B2B purchase flow (антиреференс: Энтеро).")
    bullet("«Тонкие» карточки на listing без ключевых атрибутов.")
    bullet("Копирование тяжёлых федеральных filter panels без адаптации под OEM-каталог ЗПМ.")
    bullet("Игнорирование региональных пробелов — Barnaul gap, Far East PARTIAL coverage.")
    bullet("Массовые маркетплейсы и агрегаторы (Tier D/E) как UX-референс для manufacturer site.")

    # 12. Приоритеты
    heading("12. Приоритеты разработки")
    para("На основании Master Report Section 9 и W3Y readiness:")
    bullet("P1 cohort для W4 intelligence: Tier A OEM + Trapeza + W1D federal anchors.")
    bullet("Приоритет UX deep-dive: filter UX (УЗНМ), view switchers (КЛЕН), information density (Trapeza), listing cards (Юниторг).")
    bullet("Региональная архитектура: multi-city coverage для Сибири/Урала/ДВ.")
    bullet("Expansion triage: Strong candidates (Kobor, НОТИС, Практика) — SERP-visible но не в Approved registry.")
    bullet("Resolve 13 Review Required tier flags — parallel, non-blocking.")
    bullet("Завершение operator manual review для non-highlighted entities — optional, not W4 gate.")
    commentary(
        "В первую очередь — каталог и карточки (P1 cohort + UX-паттерны benchmark group), "
        "потому что это ядро B2B-сайта производителя. Региональная архитектура и expansion triage — "
        "следующий слой после базовой структуры каталога."
    )

    # 13. Практические выводы
    heading("13. Практические выводы")
    bullet("Исследование создало воспроизводимый baseline: 126 entities, 46 Approved, полная tier/geography taxonomy.")
    bullet("SERP map подтверждает: для SEO-стратегии ЗПМ критичны федеральные peers (Trapeza, КЛЕН) и региональные visible players.")
    bullet("W3Y operator insights дают конкретные UX entry points — не абстрактные рекомендации, а зафиксированные patterns.")
    bullet("Benchmark group — attention list, не финальный дизайн-гайд; требует W4 deep intelligence validation.")
    bullet("Executive package позволяет через год восстановить контекст исследования без доступа к рабочим файлам.")

    # 14. Итог
    heading("14. Итог исследования")
    para(
        "Программа BZPM Market Intelligence (W1–W3Y) завершила фазу discovery, registry consolidation "
        "и operator insight capture. Сформирована полная карта конкурентного поля с 46 Approved entities, "
        "43 expansion candidates и зафиксированной SERP visibility map. Операторские UX-наблюдения "
        "определяют направление для разработки нового сайта ЗПМ: information density, simplified filters, "
        "view switchers, structured listing cards и regional architecture."
    )
    para(
        "W4 Competitor Intelligence может начинаться на approved P1 cohort с operator observations "
        "как routing context. Настоящий Executive Package служит долгосрочным аналитическим продуктом "
        "для клиента, руководства и новых участников проекта."
    )
    commentary(
        "Используйте BZPM Market Research.xlsx для обзора данных и диаграмм, "
        "этот документ — для выводов и приоритетов. Authority markdown в родительской папке — "
        "для верификации и углублённой работы. Версия v2.1 RU: русская локализация и клиентские пояснения без изменения данных."
    )

    out_path = OUT / "BZPM Research Conclusions.docx"
    doc.save(out_path)
    return out_path


def main():
    entities = load_registry()
    assert len(entities) == 126, f"Expected 126 entities, got {len(entities)}"
    charts = build_excel(entities)
    word_path = build_word()
    print("Excel:", OUT / "BZPM Market Research.xlsx")
    print("Word:", word_path)
    print("Charts:", len(charts))
    for c in charts:
        print(" -", c)


if __name__ == "__main__":
    main()
