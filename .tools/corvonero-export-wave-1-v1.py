#!/usr/bin/env python3
"""Corvonero Export Wave 1 — LP-01 DOCX + Consolidated Research XLSX."""

from __future__ import annotations

import hashlib
import json
import os
import re
import zipfile
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

REPO = Path(r"X:\AI MARS")
STORAGE_ROOT = Path(r"X:\AI MARS STORAGE")
OUT = STORAGE_ROOT / "exports/corvonero/CORVONERO-EXPORT-WAVE-1-2026-06-29"
CHECKPOINT = "2de6bafab4ca80f2e1bf641468f0b973c4c21282"
EXPORT_DATE = "2026-06-29"

PILOTS = REPO / "projects/mars-search-ppc-production/pilots/corvonero"
ORCA = REPO / "projects/orca/projects/corvonero-direct-v2-clean-room"
MIG = REPO / "incoming/mig/pilots/corvonero/session-mig-20260622-corv01"

DOCX_NAME = "CORVONERO-LP01-ТЕКСТ-ДЛЯ-РОМАНА-v1.docx"
XLSX_NAME = "CORVONERO-СВОДНОЕ-ИССЛЕДОВАНИЕ-v1.xlsx"


def load_json(path: Path):
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def rel(p: Path) -> str:
    try:
        return str(p.relative_to(REPO)).replace("\\", "/")
    except ValueError:
        return str(p)


# --- DOCX helpers ---

def add_label_paragraph(doc, label: str, value: str | None = None):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label} ")
    run_label.bold = True
    if value:
        p.add_run(value)


def add_bullet_list(doc, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_lp01_docx(path: Path) -> None:
    faq = load_json(PILOTS / "CORVONERO-PHASE-6.6-LP01-FINAL-FAQ-v3.json")
    faq_answers = {
        "faq-01": "Программист 1С в «Корво Неро» берёт практические задачи по вашей базе: доработка конфигураций, исправление ошибок, настройка учёта и бизнес-процессов, отчёты и обработки, интеграции с сайтами и Битрикс, маркировка и Честный знак, обновление и сопровождение баз, разовые и срочные работы.",
        "faq-02": "Работаем с конфигурациями: УТ, УНФ, Розница, КА, БП (Бухгалтерия предприятия). Уточните версию и редакцию вашей базы при обращении.",
        "faq-03": "Да. Удалённое подключение к базе доступно по всей России. Формат согласуется при постановке задачи.",
        "faq-04": "Да, выезд специалиста возможен в пределах Новосибирска. С клиентами из других городов работаем удалённо.",
        "faq-05": "Минимальный заказ — 2 часа.",
        "faq-06": "Стоимость работы — от 3 000 ₽ в час. Итог зависит от сложности задачи, конфигурации, объёма доработок и необходимости выезда в Новосибирск. Перед началом работ согласуем оценку по описанию задачи.",
        "faq-07": "Да. Работаем с юридическими лицами и ИП по договору с безналичной оплатой.",
        "faq-08": "Да, возможность срочного подключения зависит от текущей загрузки специалиста. Опишите проблему — мы уточним доступный формат работы.",
        "faq-09": "При первом разговоре уточним конфигурацию и версию 1С, что именно нужно сделать или какая ошибка возникла. Для заявки на сайте достаточно оставить номер телефона.",
    }

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate([
        "Корво Неро",
        "Посадочная страница «Программист 1С»",
        "Текст для сборки страницы",
    ]):
        run = title.add_run(line if i == 0 else "\n" + line)
        run.bold = True
        run.font.size = Pt(14 if i == 0 else 12)

    doc.add_paragraph()

    def block_heading(n: int, name: str):
        h = doc.add_heading(f"БЛОК {n}. {name}", level=1)

    # Block 1
    block_heading(1, "Первый экран")
    add_label_paragraph(doc, "Заголовок:", "Программист 1С для доработки, настройки и исправления ошибок")
    add_label_paragraph(doc, "Текст:", "Решаем разовые задачи в 1С, дорабатываем конфигурации и восстанавливаем работу базы. Удалённо по России, с выездом в Новосибирске.")
    add_label_paragraph(doc, "Телефон:", "+7 (383) 390-29-28")
    add_label_paragraph(doc, "Кнопка:", "Обсудить задачу")
    add_label_paragraph(doc, "Кнопка:", "Получить оценку")

    # Block 2
    block_heading(2, "Для кого услуга")
    add_label_paragraph(doc, "Заголовок:", "Услуги программиста 1С для компаний и ИП")
    add_label_paragraph(
        doc,
        "Текст:",
        "Подключаемся к разовым и текущим задачам: дорабатываем конфигурации, исправляем ошибки, настраиваем учёт и интеграции. Подходит компаниям, которым нужен внешний специалист без найма программиста в штат.\n\n"
        "Обращаются руководители, бухгалтерия и IT-ответственные, у которых уже используется 1С: УТ, УНФ, Розница, КА или Бухгалтерия предприятия.",
    )

    # Block 3
    block_heading(3, "Что делает программист 1С")
    add_label_paragraph(doc, "Заголовок:", "Что делает программист 1С")
    add_label_paragraph(doc, "Текст:", "Специалист 1С в «Корво Неро» выполняет практические задачи по вашей базе:")
    add_bullet_list(doc, [
        "доработка конфигураций 1С под процессы компании",
        "исправление ошибок и восстановление работоспособности",
        "настройка учёта и бизнес-процессов",
        "разработка отчётов, обработок и печатных форм",
        "интеграции с сайтами, Битрикс и внешними системами",
        "настройка маркировки и работы с Честным знаком",
        "обновление и сопровождение баз",
        "разовые и срочные задачи",
    ])
    add_label_paragraph(doc, "Кнопка:", "Получить оценку")

    # Block 4
    block_heading(4, "Типовые задачи")
    add_label_paragraph(doc, "Заголовок:", "Типовые задачи")
    add_label_paragraph(doc, "Текст:", "Частые запросы, с которыми обращаются к программисту 1С:")
    add_bullet_list(doc, [
        "база перестала проводить документы или выдаёт ошибку при закрытии периода",
        "нужен нестандартный отчёт или печатная форма",
        "требуется обмен данными между 1С и интернет-магазином или CRM",
        "нужно настроить маркировку товаров в конфигурации",
        "после обновления перестали работать ранее выполненные доработки",
        "нужна разовая помощь специалиста без долгосрочного сопровождения",
    ])
    add_label_paragraph(doc, "Текст:", "Опишите задачу — мы уточним объём работ и предложим формат подключения.")
    add_label_paragraph(doc, "Кнопка:", "Обсудить задачу")

    # Block 5
    block_heading(5, "Конфигурации 1С")
    add_label_paragraph(doc, "Заголовок:", "Конфигурации 1С")
    add_label_paragraph(doc, "Текст:", "Работаем с типовыми конфигурациями:")
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    headers = [("Конфигурация", "Кратко"), ("1С:УТ", "торговля и склад"), ("1С:УНФ", "управление небольшой фирмой"),
               ("1С:Розница", "розничные продажи"), ("1С:КА", "комплексная автоматизация"), ("1С:БП", "бухгалтерия предприятия")]
    for i, (a, b) in enumerate(headers):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        if i == 0:
            for cell in table.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
    add_label_paragraph(doc, "Текст:", "Если ваша база построена на одной из этих конфигураций — опишите версию и суть задачи при обращении.")
    add_label_paragraph(doc, "Кнопка:", "Получить оценку")

    # Block 6
    block_heading(6, "Формат работы")
    add_label_paragraph(doc, "Заголовок:", "Формат работы")
    add_label_paragraph(doc, "Текст:", "Удалённо — работаем с клиентами по всей России: подключаемся к базе, обсуждаем задачу и передаём результат дистанционно.")
    add_label_paragraph(doc, "Текст:", "На объекте — выезд специалиста возможен в пределах Новосибирска. Задачи из других городов выполняем удалённо.")
    add_label_paragraph(doc, "Кнопка:", "Заказать звонок")

    # Block 7
    block_heading(7, "Стоимость")
    add_label_paragraph(doc, "Заголовок:", "Стоимость работы программиста 1С")
    add_label_paragraph(doc, "Текст:", "Стоимость работы — от 3 000 ₽ в час.")
    add_label_paragraph(doc, "Текст:", "Минимальный заказ — 2 часа.")
    add_label_paragraph(
        doc,
        "Текст:",
        "Итоговая сумма зависит от сложности задачи, конфигурации, объёма доработок и необходимости выезда в Новосибирск. Перед началом работ согласуем оценку по вашему описанию.",
    )
    add_label_paragraph(doc, "Кнопка:", "Получить оценку")

    # Block 8
    block_heading(8, "Как мы работаем")
    add_label_paragraph(doc, "Заголовок:", "Как мы работаем")
    add_bullet_list(doc, [
        "Обсуждаем задачу — по телефону, в мессенджере или через форму.",
        "Уточняем конфигурацию и объём работ — версия базы, симптомы ошибки или требования к доработке.",
        "Согласовываем оценку — ориентир по времени и стоимости в рамках почасовой модели.",
        "Выполняем работу — удалённо или с выездом в Новосибирске.",
        "Передаём результат — при необходимости кратко фиксируем выполненные изменения.",
    ])
    add_label_paragraph(doc, "Кнопка:", "Заказать звонок")

    # Block 9
    block_heading(9, "Почему обращаются в Корво Неро")
    add_label_paragraph(doc, "Заголовок:", "Почему обращаются в Корво Неро")
    add_label_paragraph(doc, "Текст:", "«Корво Неро» помогает бизнесу настраивать, дорабатывать, обновлять и сопровождать 1С.")
    add_label_paragraph(
        doc,
        "Текст:",
        "К программисту 1С обращаются, когда нужен практический результат без найма штатного разработчика: исправить ошибку, доработать конфигурацию, настроить учёт или связать 1С с другими системами.",
    )
    add_label_paragraph(doc, "Текст:", "Формат сотрудничества:")
    add_bullet_list(doc, [
        "разовые задачи — подключаемся к конкретной проблеме или доработке",
        "понятная почасовая модель — от 3 000 ₽ в час, минимальный заказ 2 часа",
        "работа по договору с безналичной оплатой для юридических лиц и ИП",
        "удалённое подключение по всей России",
        "выезд специалиста в Новосибирске",
        "работаем с конфигурациями УТ, УНФ, Розница, КА и БП",
    ])
    add_label_paragraph(doc, "Кнопка:", "Обсудить задачу")

    # Block 10 FAQ
    block_heading(10, "Частые вопросы")
    add_label_paragraph(doc, "Заголовок:", "Частые вопросы")
    for item in faq["items"]:
        q = doc.add_paragraph()
        q_run = q.add_run(item["question_ru"])
        q_run.bold = True
        ans = faq_answers.get(item["id"], item.get("answer_ru", ""))
        doc.add_paragraph(ans)

    # Block 11
    block_heading(11, "Контакты и мессенджеры")
    add_label_paragraph(doc, "Заголовок:", "Связаться с нами")
    add_label_paragraph(doc, "Текст:", "Позвоните нам или выберите удобный мессенджер.")
    add_label_paragraph(doc, "Телефон:", "+7 (383) 390-29-28")
    add_label_paragraph(doc, "Список:", "MAX, Telegram, WhatsApp")

    # Block 12
    block_heading(12, "Форма")
    add_label_paragraph(doc, "Заголовок:", "Оставить заявку на услуги программиста 1С")
    add_label_paragraph(doc, "Подпись:", "Оставьте телефон — уточним задачу и сориентируем по стоимости.")
    add_label_paragraph(doc, "Поле:", "Имя")
    add_label_paragraph(doc, "Поле:", "Телефон")
    add_label_paragraph(doc, "Кнопка:", "Заказать звонок")

    # Block 13
    block_heading(13, "Финальный призыв")
    add_label_paragraph(doc, "Заголовок:", "Обсудим вашу задачу в 1С")
    add_label_paragraph(doc, "Текст:", "Расскажите, что нужно исправить, настроить или доработать. Уточним детали и сориентируем по стоимости.")
    add_label_paragraph(doc, "Кнопка:", "Обсудить задачу")
    add_label_paragraph(doc, "Телефон:", "+7 (383) 390-29-28")

    # Block 14
    block_heading(14, "Футер")
    add_label_paragraph(doc, "Текст:", "Центр автоматизации «Корво Неро»")
    add_label_paragraph(doc, "Текст:", "ИП Никифоров Роман Вадимович")
    add_label_paragraph(doc, "Текст:", "+7 (383) 390-29-28")

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = f"Версия 1 · {EXPORT_DATE.replace('-', '.')}"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(path)


# --- XLSX helpers ---

HEADER_FONT = Font(bold=True)
HEADER_FILL = PatternFill("solid", fgColor="E8EEF4")
WRAP = Alignment(wrap_text=True, vertical="top")


def write_sheet(ws, headers: list[str], rows: list[list], freeze: bool = True):
    ws.append(headers)
    for col in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = WRAP
    for row in rows:
        ws.append(row)
    if freeze:
        ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for col_idx, header in enumerate(headers, 1):
        max_len = len(str(header))
        for r in rows[:200]:
            if col_idx - 1 < len(r):
                max_len = max(max_len, min(60, len(str(r[col_idx - 1] or ""))))
        ws.column_dimensions[get_column_letter(col_idx)].width = min(48, max(12, max_len + 2))


def domain_from_url(url: str) -> str:
    if not url:
        return ""
    if url.startswith("tel:"):
        return "tel"
    try:
        host = urlparse(url).netloc or urlparse(url).path
        host = host.lower().replace("www.", "")
        if "yabs.yandex" in host:
            return ""
        return host.split("/")[0]
    except Exception:
        return ""


def build_workbook(path: Path) -> dict:
    canonical = load_json(ORCA / "semantic-core/corvonero-canonical-phrase-registry-v1.json")
    normalized = load_json(ORCA / "semantic-core/corvonero-normalized-corpus-v1.json")
    wordstat = load_json(ORCA / "mig-source/mig-wordstat-source-ledger-v1.json")
    reviewed = {r["phrase_id"]: r for r in load_json(PILOTS / "CORVONERO-RUN-004-PHASE-5.2-FINAL-REVIEWED-REGISTRY-v1.json")["records"]}
    accept = {r["phrase_id"]: r for r in load_json(PILOTS / "CORVONERO-RUN-004-PHASE-5.2-FINAL-ACCEPT-v1.json")["records"]}
    reject = {r["phrase_id"]: r for r in load_json(PILOTS / "CORVONERO-RUN-004-PHASE-5.2-FINAL-REJECT-v1.json")["records"]}
    abstain = {r["phrase_id"]: r for r in load_json(PILOTS / "CORVONERO-RUN-004-PHASE-5.2-FINAL-ABSTAIN-v1.json")["records"]}
    backlog = {r["phrase_id"]: r for r in load_json(PILOTS / "CORVONERO-RUN-004-PHASE-4-UNPROCESSED-IDS-MANIFEST-v1.json")["records"]}
    campaigns = load_json(PILOTS / "CORVONERO-PHASE-6.1-CAMPAIGN-FAMILIES-v2.json")
    ad_groups = load_json(PILOTS / "CORVONERO-PHASE-6.1-AD-GROUP-ARCHITECTURE-v2.json")
    allocation = load_json(PILOTS / "CORVONERO-PHASE-6.1-PHRASE-ALLOCATION-v2.json")["allocations"]
    group_lp = load_json(PILOTS / "CORVONERO-PHASE-6.2-GROUP-TO-LP-MAP-v1.json")
    serp_index = load_json(MIG / "serp_r1_index.json")
    lp_audit = load_json(PILOTS / "CORVONERO-PHASE-6.1-LANDING-PAGE-AUDIT-v1.json")
    lp_matrix = load_json(PILOTS / "CORVONERO-PHASE-6.1-LANDING-PAGE-MATRIX-v2.json")
    intake = load_json(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json")
    exclusion = load_json(PILOTS / "CORVONERO-RUN-004-PHASE-5.2-FINAL-EXCLUSION-TAXONOMY-v1.json")
    export_matrix = load_json(PILOTS / "CORVONERO-EXPORT-READINESS-MATRIX-v1.json")
    campaign_ready = load_json(PILOTS / "CORVONERO-PHASE-6.2-CAMPAIGN-READINESS-MATRIX-v2.json")

    campaign_names = {c["campaign_id"]: c["working_name"] for c in campaigns["campaign_families"]}
    group_meta = {g["group_id"]: g for g in ad_groups["ad_groups"]}
    group_lp_map = {g["group_id"]: g for g in group_lp["groups"]}

    # Build LED -> canonical map
    led_to_canonical: dict[str, str] = {}
    for p in canonical["phrases"]:
        for led in p.get("source_row_ids", []):
            led_to_canonical[led] = p["phrase_id"]

    norm_by_led = {r["ledger_row_id"]: r for r in normalized["rows"]}

    wb = Workbook()
    wb.remove(wb.active)

    counts = {
        "canonical": len(canonical["phrases"]),
        "assessed": len(reviewed),
        "accept": len(accept),
        "reject": len(reject),
        "abstain": len(abstain),
        "backlog": len(backlog),
    }

    # 00_README
    ws = wb.create_sheet("00_README")
    readme_rows = [
        ["Project", "Corvonero — Яндекс Директ / LP production"],
        ["Date", EXPORT_DATE],
        ["Purpose", "Consolidated research, semantic adjudication and campaign planning evidence"],
        ["Checkpoint", CHECKPOINT],
        ["", ""],
        ["PARTIAL COVERAGE", ""],
        ["Canonical corpus", counts["canonical"]],
        ["Assessed", counts["assessed"]],
        ["Unprocessed", counts["backlog"]],
        ["Semantic coverage", "67.5%"],
        ["SERP evidence (official boundary)", "5 / 10 planned queries"],
        ["", ""],
        ["Sheet index", "Description"],
        ["01_PROJECT_CONTEXT", "Confirmed business facts"],
        ["02_SOURCE_REGISTRY", "Source artefact inventory"],
        ["03_WORDSTAT_RAW", "Wordstat source ledger rows"],
        ["04_WORDSTAT_NORMALIZED", "Normalized corpus"],
        ["05_CANONICAL_SEMANTICS", "Full canonical registry + verdicts"],
        ["06_ACCEPT", "935 ACCEPT phrases"],
        ["07_REJECT", "368 REJECT phrases"],
        ["08_ABSTAIN", "296 ABSTAIN phrases"],
        ["09_UNPROCESSED_BACKLOG", "769 unprocessed records"],
        ["10_CAMPAIGNS", "6 campaign families"],
        ["11_AD_GROUPS", "21 ad groups"],
        ["12_PHRASE_ALLOCATION", "ACCEPT phrase allocation"],
        ["13_SERP_QUERY_REGISTER", "10 planned SERP queries"],
        ["14_SERP_RESULTS", "Parsed SERP observations"],
        ["15_COMPETITORS", "Domain summary from SERP"],
        ["16_SEARCH_OBSERVATIONS", "Search market observations"],
        ["17_LANDING_PAGE_AUDIT", "Audited pages"],
        ["18_LANDING_PAGE_PROGRAM", "LP-01..LP-06 program"],
        ["19_LP01_FINAL_CONTENT_SUMMARY", "LP-01 approved fields"],
        ["20_EXCLUSION_EVIDENCE", "Reject exclusion families"],
        ["21_RISKS_AND_LIMITATIONS", "Material limitations"],
        ["22_DECISION_LOG", "Operator decisions"],
        ["23_EXPORT_READINESS", "Deliverable readiness"],
        ["", ""],
        ["Distinction", "Raw evidence = Wordstat/SERP files; Decisions = Phase 5.2 verdicts; Planning = Phase 6 architecture"],
    ]
    for r in readme_rows:
        ws.append(r)
    ws.column_dimensions["A"].width = 36
    ws.column_dimensions["B"].width = 72

    # 01_PROJECT_CONTEXT
    ctx_rows = [
        ["Brand", intake["company"]["commercial_name"], "CONFIRMED", rel(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json"), ""],
        ["Legal entity", intake["company"]["legal_entity"], "CONFIRMED", "ATLAS LE-0006 (E0 partial)", ""],
        ["Target customers", ", ".join(intake["audience"]), "CONFIRMED", rel(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json"), "B2B companies and IP"],
        ["Services", "1C programmer, support, development, integrations, marking", "CONFIRMED", "Business intake + semantic families", ""],
        ["Work formats", "Remote Russia; on-site Novosibirsk", "CONFIRMED", rel(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json"), ""],
        ["Contract", "Yes", "CONFIRMED", rel(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json"), ""],
        ["Cashless payment", "Yes", "CONFIRMED", rel(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json"), ""],
        ["Geography (launch)", ", ".join(intake["geography"]["first_launch"]), "CONFIRMED", rel(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json"), ""],
        ["Configurations", ", ".join(intake["confirmed_configurations"]), "CONFIRMED", rel(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json"), ""],
        ["Hourly rate", "3000 RUB", "CONFIRMED", rel(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json"), ""],
        ["Minimum order", "2 hours / 6000 RUB", "CONFIRMED", rel(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json"), ""],
        ["VAT", "UNKNOWN", "UNKNOWN", "intake safe_unknown", "Not confirmed"],
        ["Target CPL", "UNKNOWN", "UNKNOWN", "intake safe_unknown", ""],
        ["Official 1C partner", "NOT CONFIRMED", "PROHIBITED", "intake prohibited_claims", ""],
        ["Monthly budget", "100000 RUB", "CONFIRMED", rel(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json"), "No historical PPC data"],
    ]
    write_sheet(wb.create_sheet("01_PROJECT_CONTEXT"), ["Topic", "Value", "Status", "Source", "Notes"], ctx_rows)

    # 02_SOURCE_REGISTRY
    sources = [
        ["SRC-INTAKE", "Business Intake", rel(ORCA / "intake/corvonero-direct-v2-business-intake-v1.json"), "2026-06-22", "Company facts", "AUTHORITY", "01", ""],
        ["SRC-ATLAS", "ATLAS register", "projects/atlas/population/ATLAS-CORVONERO-ORGANIZATION-REGISTER-v1.md", "2026-06", "LE-0006 E0 partial", "AUTHORITY", "01,19", ""],
        ["SRC-WS-LEDGER", "Wordstat", rel(ORCA / "mig-source/mig-wordstat-source-ledger-v1.json"), "2026-06-22", f"{wordstat['stats']['rows_read']} rows", "RAW", "03,04", ""],
        ["SRC-WS-NORM", "Wordstat normalized", rel(ORCA / "semantic-core/corvonero-normalized-corpus-v1.json"), "2026-06-22", f"{normalized['row_count']} rows", "DERIVED", "04,05", ""],
        ["SRC-CANONICAL", "Canonical corpus", rel(ORCA / "semantic-core/corvonero-canonical-phrase-registry-v1.json"), "2026-06-22", "2368 phrases", "AUTHORITY", "05-09", ""],
        ["SRC-ADJUDICATION", "Phase 5.2 verdicts", rel(PILOTS / "CORVONERO-RUN-004-PHASE-5.2-FINAL-REVIEWED-REGISTRY-v1.json"), "2026-06-28", "1599 assessed", "AUTHORITY", "05-08", ""],
        ["SRC-CAMPAIGN", "Campaign architecture", rel(PILOTS / "CORVONERO-PHASE-6.1-CAMPAIGN-FAMILIES-v2.json"), "2026-06-28", "6 campaigns", "PLANNING", "10-12", ""],
        ["SRC-SERP", "SERP captures", rel(MIG / "serp_r1_index.json"), "2026-06-22", "5/10 official; 7 Grade B files exist", "RAW_PARTIAL", "13-15", ""],
        ["SRC-LP-AUDIT", "LP audit", rel(PILOTS / "CORVONERO-PHASE-6.1-LANDING-PAGE-AUDIT-v1.json"), "2026-06-28", "2 sites", "EVIDENCE", "17", ""],
        ["SRC-LP01-COPY", "LP-01 final copy v3", rel(PILOTS / "CORVONERO-PHASE-6.6-LP01-FINAL-PRODUCTION-COPY-v3.json"), "2026-06-29", "Operator approved", "AUTHORITY", "19,D2", ""],
    ]
    write_sheet(wb.create_sheet("02_SOURCE_REGISTRY"), ["Source ID", "Source type", "File / artefact", "Date", "Coverage", "Authority level", "Used in sheets", "Notes"], sources)

    # 03_WORDSTAT_RAW
    ws_rows = []
    for r in wordstat["rows"]:
        ws_rows.append([
            r.get("source_file"), r.get("source_sheet"), r.get("source_row"), r.get("original_phrase"),
            r.get("original_frequency"), r.get("region_context"), r.get("mig_source_id"), "2026-06-22",
            r.get("original_frequency"), r.get("provenance", ""),
        ])
    write_sheet(wb.create_sheet("03_WORDSTAT_RAW"), [
        "Source file", "Source sheet", "Original row", "Original query", "Frequency", "Region",
        "Collection pass", "Collection date", "Raw value", "Notes",
    ], ws_rows)

    # 04_WORDSTAT_NORMALIZED
    norm_rows = []
    for r in normalized["rows"]:
        cid = led_to_canonical.get(r["ledger_row_id"], "")
        norm_rows.append([
            r["ledger_row_id"], r.get("original_phrase"), r.get("normalized"), r.get("original_frequency"),
            r.get("region_context"), r.get("source_file"), "NORMALIZED", r.get("dedupKey"), "INCLUDED" if cid else "EXCLUDED", "" if cid else "Not in canonical",
        ])
    write_sheet(wb.create_sheet("04_WORDSTAT_NORMALIZED"), [
        "Normalized row ID", "Original phrase", "Normalized phrase", "Frequency", "Region", "Source file",
        "Normalization status", "Duplicate cluster", "Canonical inclusion status", "Exclusion reason",
    ], norm_rows)

    # 05_CANONICAL_SEMANTICS
    sem_rows = []
    for p in canonical["phrases"]:
        pid = p["phrase_id"]
        rev = reviewed.get(pid)
        if rev:
            assessment = "ASSESSED"
            verdict = rev.get("phase52_final_verdict") or rev.get("operator_final_verdict") or ""
            reason = rev.get("phase52_rationale") or rev.get("rationale") or ""
            intent = rev.get("primary_intent")
            sf = rev.get("service_family")
            geo = (rev.get("geography") or {}).get("status")
            flags = rev.get("review_flag_root_cause")
            run_id = "corv-semantic-v2-20260626-004"
        elif pid in backlog:
            assessment = "UNPROCESSED_BACKLOG"
            verdict = ""
            reason = backlog[pid].get("reason", "Not assessed — Phase 4 endpoint failure")
            intent = sf = geo = flags = ""
            run_id = ""
        else:
            assessment = "UNPROCESSED_BACKLOG"
            verdict = ""
            reason = "Not in reviewed registry"
            intent = sf = geo = flags = ""
            run_id = ""
        dup = p.get("duplicate_count", 0)
        cluster = p["phrase_id"] if dup == 0 else f"cluster-{p.get('source_row_ids', [''])[0]}"
        sem_rows.append([
            pid, p.get("phrase"), p.get("normalized_phrase"), p.get("combined_frequency"),
            ";".join(p.get("provenance", [])), cluster, "CANONICAL", assessment, verdict, reason,
            intent, sf, geo, flags, run_id,
        ])
    write_sheet(wb.create_sheet("05_CANONICAL_SEMANTICS"), [
        "Canonical ID", "Original phrase", "Normalized phrase", "Frequency", "Source", "Duplicate cluster",
        "Canonical status", "Assessment status", "Final verdict", "Verdict reason", "Intent", "Service family",
        "Geo metadata", "Review flags", "Run ID",
    ], sem_rows)

    # 06_ACCEPT
    acc_rows = []
    for pid, r in sorted(accept.items()):
        alloc = allocation.get(pid, {})
        cid = alloc.get("campaign_id", "")
        gid = alloc.get("ad_group_id", "")
        glp = group_lp_map.get(gid, {})
        acc_rows.append([
            pid, r.get("phrase"), (r.get("source_metadata") or {}).get("combined_frequency"),
            cid, campaign_names.get(cid, ""), gid, group_meta.get(gid, {}).get("working_name", ""),
            r.get("primary_intent"), r.get("secondary_intent"), campaigns["campaign_families"][0].get("default_priority") if cid == "CA-01" else "P1",
            glp.get("landing_page_id", ""), glp.get("landing_page_id", "").replace("LP-", "LP-0") if False else (
                {"LP-01": "Программист / специалист 1С", "LP-02": "Сопровождение 1С", "LP-03": "Доработка 1С",
                 "LP-04": "Интеграции 1С", "LP-05": "Маркировка / Честный знак", "LP-06": "Отчёты и обработки 1С"}.get(glp.get("landing_page_id", ""), "")
            ),
            (r.get("geography") or {}).get("status"), "",
        ])
    write_sheet(wb.create_sheet("06_ACCEPT"), [
        "Canonical ID", "Phrase", "Frequency", "Campaign ID", "Campaign name", "Ad-group ID", "Ad-group name",
        "Primary intent", "Secondary intent", "Priority", "LP ID", "LP name", "Geography metadata", "Notes",
    ], acc_rows)

    lp_names = {"LP-01": "Программист / специалист 1С", "LP-02": "Сопровождение 1С", "LP-03": "Доработка 1С",
                "LP-04": "Интеграции 1С", "LP-05": "Маркировка / Честный знак", "LP-06": "Отчёты и обработки 1С"}

    # Fix accept rows LP names properly
    ws_acc = wb["06_ACCEPT"]
    for row_idx in range(2, ws_acc.max_row + 1):
        gid = ws_acc.cell(row_idx, 6).value
        glp = group_lp_map.get(gid, {})
        lp_id = glp.get("landing_page_id", "")
        ws_acc.cell(row_idx, 11, lp_id)
        ws_acc.cell(row_idx, 12, lp_names.get(lp_id, ""))

    # 07_REJECT
    rej_rows = []
    for pid, r in sorted(reject.items()):
        rej_rows.append([
            pid, r.get("phrase"), (r.get("source_metadata") or {}).get("combined_frequency"),
            r.get("exclusion_family"), r.get("phase52_rationale") or r.get("rationale"), r.get("exclusion_family"),
            "", "LOW" if r.get("exclusion_family") else "MEDIUM", "",
        ])
    write_sheet(wb.create_sheet("07_REJECT"), [
        "Canonical ID", "Phrase", "Frequency", "Reject category", "Reason", "Exclusion family",
        "Potential minus-word evidence", "Risk of overblocking", "Notes",
    ], rej_rows)

    # 08_ABSTAIN
    abs_rows = []
    for pid, r in sorted(abstain.items()):
        abs_rows.append([
            pid, r.get("phrase"), (r.get("source_metadata") or {}).get("combined_frequency"),
            r.get("primary_intent"), r.get("phase52_rationale") or r.get("rationale"), "Manual review if backlog processed", "",
        ])
    write_sheet(wb.create_sheet("08_ABSTAIN"), [
        "Canonical ID", "Phrase", "Frequency", "Abstain category", "Reason", "Potential future action", "Notes",
    ], abs_rows)

    # 09_UNPROCESSED_BACKLOG
    bl_rows = [[pid, r.get("phrase"), r.get("combined_frequency"), "Wordstat canonical", "UNPROCESSED_BACKLOG",
                r.get("reason", "Phase 4 endpoint failure"), "Requires semantic run continuation"] for pid, r in sorted(backlog.items())]
    write_sheet(wb.create_sheet("09_UNPROCESSED_BACKLOG"), [
        "Canonical ID", "Phrase", "Frequency", "Source", "Status", "Reason not assessed", "Future processing requirement",
    ], bl_rows)

    # 10_CAMPAIGNS
    cr = {c["campaign_id"]: c for c in campaign_ready["campaigns"]}
    camp_rows = []
    for c in campaigns["campaign_families"]:
        cid = c["campaign_id"]
        camp_rows.append([
            cid, c["working_name"], c.get("default_priority"), c.get("allocated_phrase_count"),
            "Search acquisition for 1C service family", ", ".join(c.get("primary_intents", [])),
            c.get("landing_page", {}).get("required", ""), cr.get(cid, {}).get("final_state", ""),
            cr.get(cid, {}).get("blocker", ""), c.get("launch_geography", ""), "Not in first launch scope",
        ])
    write_sheet(wb.create_sheet("10_CAMPAIGNS"), [
        "Campaign ID", "Campaign name", "Priority", "Phrase count", "Purpose", "Primary intents", "Landing page",
        "Current readiness", "Blockers", "Launch geography", "Expansion status",
    ], camp_rows)

    # 11_AD_GROUPS
    ag_rows = []
    for g in ad_groups["ad_groups"]:
        glp = group_lp_map.get(g["group_id"], {})
        ag_rows.append([
            g["campaign_id"], campaign_names.get(g["campaign_id"], ""), g["group_id"], g.get("working_name"),
            len(g.get("phrase_ids", [])), g.get("primary_intent"), "", "", glp.get("landing_page_id", ""),
            "P1", "REQUIREMENTS_ONLY" if glp.get("landing_page_id") != "LP-01" else "LP01_FINAL_COPY_READY", "",
        ])
    write_sheet(wb.create_sheet("11_AD_GROUPS"), [
        "Campaign ID", "Campaign name", "Ad-group ID", "Ad-group name", "Phrase count", "Primary intent",
        "Permitted secondary intent", "Prohibited intent", "LP ID", "Priority", "Readiness", "Notes",
    ], ag_rows)

    # 12_PHRASE_ALLOCATION
    alloc_rows = []
    for pid in sorted(allocation.keys()):
        a = allocation[pid]
        cid = a["campaign_id"]
        gid = a["ad_group_id"]
        glp = group_lp_map.get(gid, {})
        alloc_rows.append([
            pid, accept.get(pid, {}).get("phrase", ""), cid, campaign_names.get(cid, ""), gid,
            group_meta.get(gid, {}).get("working_name", ""), glp.get("landing_page_id", ""), "ALLOCATED", "PASS",
        ])
    write_sheet(wb.create_sheet("12_PHRASE_ALLOCATION"), [
        "Canonical ID", "Phrase", "Campaign ID", "Campaign name", "Ad-group ID", "Ad-group name", "LP ID",
        "Allocation status", "Duplicate check",
    ], alloc_rows)

    # 13_SERP_QUERY_REGISTER — official 5/10 boundary
    official_completed = {"r1q01", "r1q02", "r1q03", "r1q04", "r1q05"}
    serp_reg_rows = []
    for q in serp_index["queries"]:
        qid = q["r1_id"]
        zpm = q.get("zpm_workflow_capture", {})
        if qid in official_completed:
            status = "COMPLETED"
        else:
            status = "INCOMPLETE/BLOCKED"
        captcha = zpm.get("captcha_status", q.get("playwright_capture", {}).get("captcha_status", ""))
        grade = zpm.get("evidence_grade", q.get("evidence_grade", ""))
        ev_path = zpm.get("serp_json", q.get("artifact_zpm_workflow", q.get("artifact", "")))
        notes = ""
        if qid in {"r1q08", "r1q10"} and zpm.get("captcha_status") == "none":
            notes = "Additional Grade B capture on disk; not counted in official 5/10 checkpoint boundary"
        serp_reg_rows.append([
            qid, q.get("query"), "Новосибирск", "mobile", status, zpm.get("captured_at", serp_index.get("collection_date")),
            grade, captcha, ev_path, notes,
        ])
    write_sheet(wb.create_sheet("13_SERP_QUERY_REGISTER"), [
        "Query ID", "Query", "Region", "Device", "Status", "Execution date", "Evidence grade", "CAPTCHA / failure",
        "Evidence paths", "Notes",
    ], serp_reg_rows)

    # 14_SERP_RESULTS
    serp_result_rows = []
    serp_base = MIG
    for q in serp_index["queries"]:
        qid = q["r1_id"]
        zpm = q.get("zpm_workflow_capture", {})
        json_rel = zpm.get("serp_json")
        if not json_rel or zpm.get("captcha_status") == "blocked":
            continue
        serp_path = serp_base / json_rel
        if not serp_path.exists():
            continue
        data = load_json(serp_path)
        pos = 0
        for item in data.get("organic_results", []):
            pos += 1
            domain = domain_from_url(item.get("url", "")) or (item.get("path_text", "").split("›")[0].strip() if item.get("path_text") else "")
            serp_result_rows.append([
                qid, pos, "organic", domain, item.get("url", ""), item.get("title", ""), item.get("path_text", ""),
                "", "", "", "", "", "", rel(serp_path), data.get("evidence_grade", zpm.get("evidence_grade", "B")),
            ])
        for item in data.get("visible_ads", []) or []:
            serp_result_rows.append([
                qid, "ad", "advertising", domain_from_url(item.get("url", "")), item.get("url", ""),
                item.get("title", ""), item.get("snippet", ""), item.get("advertiser", ""), "", "", "", "", "",
                rel(serp_path), "B",
            ])
    write_sheet(wb.create_sheet("14_SERP_RESULTS"), [
        "Query ID", "Position / section", "Result type", "Domain", "URL", "Title", "Snippet",
        "Advertiser / brand", "Commercial angle", "Service", "Price signal", "Geo signal", "Landing-page type",
        "Observed feature", "Evidence file", "Confidence",
    ], serp_result_rows)

    # 15_COMPETITORS — derived from SERP only
    domain_queries: dict[str, set] = defaultdict(set)
    domain_organic: Counter = Counter()
    domain_ads: Counter = Counter()
    domain_titles: dict[str, list] = defaultdict(list)
    for row in serp_result_rows:
        qid, _, rtype, domain = row[0], row[1], row[2], row[3]
        if not domain or domain == "tel":
            continue
        domain_queries[domain].add(qid)
        if rtype == "organic":
            domain_organic[domain] += 1
        else:
            domain_ads[domain] += 1
        if row[5]:
            domain_titles[domain].append(row[5][:120])
    comp_rows = []
    for domain in sorted(domain_queries.keys(), key=lambda d: -(domain_organic[d] + domain_ads[d])):
        comp_rows.append([
            domain, ", ".join(sorted(domain_queries[domain])), domain_organic[domain], domain_ads[domain],
            "", "; ".join(domain_titles[domain][:3]), "", "Novosibirsk (SERP context)", "",
            "", "", rel(MIG / "serp_r1_index.json"), "MEDIUM — SERP-only, identity may be incomplete",
        ])
    write_sheet(wb.create_sheet("15_COMPETITORS"), [
        "Domain / brand", "Observed queries", "Organic presence", "Advertising presence", "Services promoted",
        "Commercial claims", "Pricing claims", "Geography", "Landing-page patterns", "Strengths observed",
        "Weaknesses observed", "Evidence references", "Confidence",
    ], comp_rows[:200])

    # 16_SEARCH_OBSERVATIONS — evidence-based only
    obs_rows = [
        ["OBS-01", "Intent patterns", "Head query «программист 1С Новосибирск» mixes job boards, education and service providers", "r1q01", "Service LP must differentiate from HR/education", "MEDIUM", "Single regional SERP sample"],
        ["OBS-02", "Price language", "Commercial hourly/pricing language not dominant in captured SERP titles for programmer query", "r1q01", "LP pricing block supports differentiation", "LOW", "Partial SERP"],
        ["OBS-03", "Local vs remote", "Regional queries used; operator confirms remote Russia + onsite Novosibirsk", "intake", "Match copy to dual format", "HIGH", "Intake confirmed"],
        ["OBS-04", "Specialist vs company", "Mix of aggregators, franchises and individual specialists in SERP paths", "r1q01-r1q05", "B2B specialist positioning required", "MEDIUM", "SERP partial"],
        ["OBS-05", "Trust signals", "Existing lk.corvonero.ru has phone/brand; lacks dedicated service proposition", "LP audit", "Dedicated LP needed per CA", "HIGH", "HTTP audit"],
        ["OBS-06", "CAPTCHA impact", "SERP capture blocked for r1q06, r1q07, r1q09; limits marking/troubleshooting evidence", "serp_r1_index", "Do not infer market structure for blocked queries", "HIGH", "Documented failure"],
    ]
    write_sheet(wb.create_sheet("16_SEARCH_OBSERVATIONS"), [
        "Observation ID", "Topic", "Observation", "Query evidence", "Commercial implication", "Confidence", "Limitation",
    ], obs_rows)

    # 17_LANDING_PAGE_AUDIT
    audit_rows = [
        ["https://corvonero.ru", lp_audit["corvonero_ru_status"]["http_status"], "Corporate root", lp_audit["corvonero_ru_status"]["title"], "", "", "", "", "LP_NOT_SUITABLE", "IIS default page", "Not usable as LP", rel(PILOTS / "CORVONERO-PHASE-6.1-LANDING-PAGE-AUDIT-v1.json")],
        ["https://lk.corvonero.ru/", lp_audit["lk_corvonero_ru_status"]["http_status"], "Tilda homepage", lp_audit["lk_corvonero_ru_status"]["title"], "Корво Неро", "Generic 1C automation", "Phone CTA", "Brand, phone", "LP_GENERIC_FALLBACK", "No dedicated service URLs", "Blocks ad creation for dedicated intents", rel(PILOTS / "CORVONERO-PHASE-6.1-LANDING-PAGE-AUDIT-v1.json")],
    ]
    write_sheet(wb.create_sheet("17_LANDING_PAGE_AUDIT"), [
        "URL", "HTTP status", "Page type", "Title", "H1", "Services visible", "CTA", "Trust signals",
        "Campaign suitability", "Suitability class", "Risks", "Evidence",
    ], audit_rows)

    # 18_LANDING_PAGE_PROGRAM
    lp_prog = [
        ["LP-01", "Программист / специалист 1С", "CA-01", 404, "P1", "READY", "FINAL v3 APPROVED", "READY_FOR_TILDA_BUILD", "1", "None — final copy complete"],
        ["LP-02", "Сопровождение 1С", "CA-02", 155, "P1", "READY", "REQUIREMENTS_ONLY", "REQUIREMENTS_ONLY", "2", "Final copy not authored"],
        ["LP-03", "Доработка 1С", "CA-03", 71, "P1", "READY", "REQUIREMENTS_ONLY", "REQUIREMENTS_ONLY", "3", "Final copy not authored"],
        ["LP-04", "Интеграции 1С", "CA-04", 48, "P1", "READY", "REQUIREMENTS_ONLY", "REQUIREMENTS_ONLY", "4", "Final copy not authored"],
        ["LP-05", "Маркировка / Честный знак", "CA-05", 220, "P1", "READY", "REQUIREMENTS_ONLY", "REQUIREMENTS_ONLY", "5", "Final copy not authored"],
        ["LP-06", "Отчёты и обработки 1С", "CA-06", 37, "P2", "READY", "DEFERRED", "DEFERRED/P2", "6", "Deferred priority"],
    ]
    write_sheet(wb.create_sheet("18_LANDING_PAGE_PROGRAM"), [
        "LP ID", "Page name", "Campaign", "Phrase count", "Priority", "Requirements readiness",
        "Final-copy readiness", "Current status", "Production sequence", "Main content gaps",
    ], lp_prog)

    # 19_LP01_FINAL_CONTENT_SUMMARY
    lp01_rows = [
        ["URL", "https://lk.corvonero.ru/programmist-1s/", rel(PILOTS / "CORVONERO-PHASE-6.6-LP01-FINAL-PRODUCTION-COPY-v3.md")],
        ["Title", "Программист 1С в Новосибирске — услуги специалиста | Корво Неро", rel(PILOTS / "CORVONERO-PHASE-6.6-LP01-FINAL-PRODUCTION-COPY-v3.json")],
        ["Description", "Подключим программиста 1С для доработки...", "meta block v3"],
        ["H1", "Программист 1С для доработки, настройки и исправления ошибок", "first_screen v3"],
        ["Price", "от 3 000 ₽/час", "pricing section"],
        ["Minimum order", "2 hours", "pricing + FAQ"],
        ["Configurations", "УТ, УНФ, Розница, КА, БП", "configurations section"],
        ["Geography", "Remote Russia; onsite Novosibirsk", "work_format + FAQ"],
        ["Form fields", "Name (optional), Phone (required)", "form section"],
        ["CTAs", "Обсудить задачу; Получить оценку; Заказать звонок", "copy v3"],
        ["Messengers", "MAX, Telegram, WhatsApp", "contact section"],
        ["FAQ count", "9", rel(PILOTS / "CORVONERO-PHASE-6.6-LP01-FINAL-FAQ-v3.json")],
        ["Final checkpoint", CHECKPOINT, "corvonero-pre-export-production-2026-06"],
    ]
    write_sheet(wb.create_sheet("19_LP01_FINAL_CONTENT_SUMMARY"), ["Field", "Approved value", "Source"], lp01_rows)

    # 20_EXCLUSION_EVIDENCE
    ex_rows = []
    for fam in exclusion["families"]:
        ids = fam.get("evidence_phrase_ids", [])
        examples = []
        for eid in ids[:5]:
            if eid in reject:
                examples.append(reject[eid].get("phrase", eid))
        ex_rows.append([
            fam.get("family_id"), fam.get("family_name"), len(ids), "; ".join(examples),
            "CAMPAIGN", "MEDIUM", "DESIGN_ONLY — NOT_FINAL_MINUS_LIST",
        ])
    write_sheet(wb.create_sheet("20_EXCLUSION_EVIDENCE"), [
        "Exclusion family", "Description", "Evidence count", "Example phrases", "Recommended level",
        "Overblocking risk", "Deployment status",
    ], ex_rows)

    # 21_RISKS_AND_LIMITATIONS
    risk_rows = [
        ["R-01", "Semantics", "769 unprocessed canonical phrases (32.5% backlog)", "HIGH", "YES", "Complete semantic run", "OPEN", rel(PILOTS / "CORVONERO-RUN-004-PHASE-4-UNPROCESSED-IDS-MANIFEST-v1.json")],
        ["R-02", "Semantics", "Partial coverage 67.5% — verdicts not final for full corpus", "HIGH", "YES", "Operator sign-off only partial", "ACCEPTED_PARTIAL", rel(PILOTS / "CORVONERO-RUN-004-PHASE-5.2-PARTIAL-SEMANTIC-SIGN-OFF-v1.json")],
        ["R-03", "SERP", "5/10 planned queries in official boundary; CAPTCHA blocked others", "MEDIUM", "PARTIAL", "Throttled re-capture", "OPEN", rel(MIG / "serp_r1_index.json")],
        ["R-04", "Conversion", "No live campaign conversion data", "MEDIUM", "YES", "Launch + measure", "OPEN", "intake"],
        ["R-05", "Budget", "No validated budget forecasts from historical data", "MEDIUM", "NO", "Pilot budgeting", "OPEN", "intake"],
        ["R-06", "LP production", "LP-02..LP-06 lack final copy", "HIGH", "YES", "Sequential LP production", "OPEN", rel(PILOTS / "CORVONERO-EXPORT-READINESS-MATRIX-v1.json")],
        ["R-07", "Ads", "No authored ads", "HIGH", "YES", "Ad creation wave", "NOT_STARTED", rel(PILOTS / "CORVONERO-EXPORT-READINESS-MATRIX-v1.json")],
        ["R-08", "Commander", "No Commander import workbook", "HIGH", "YES", "D3 export after ads", "NOT_STARTED", rel(PILOTS / "CORVONERO-EXPORT-READINESS-MATRIX-v1.json")],
    ]
    write_sheet(wb.create_sheet("21_RISKS_AND_LIMITATIONS"), [
        "Risk ID", "Area", "Risk", "Severity", "Blocking", "Mitigation", "Status", "Source",
    ], risk_rows)

    # 22_DECISION_LOG
    dec_rows = [
        ["D-52-01", "Phase 5.2", "Partial semantic sign-off", "1599 assessed; 935 ACCEPT", "ACTIVE", rel(PILOTS / "CORVONERO-RUN-004-PHASE-5.2-PARTIAL-SEMANTIC-SIGN-OFF-v1.json")],
        ["D-52-02", "Phase 5.2", "47 operator adjudication overrides applied", "Selected phrase IDs", "ACTIVE", rel(PILOTS / "CORVONERO-RUN-004-PHASE-5.2-OPERATOR-DECISIONS-v1.json")],
        ["D-61-01", "Phase 6.1", "Campaign architecture consolidated v2 (6 campaigns, 21 groups)", "935 ACCEPT phrases", "ACTIVE", rel(PILOTS / "CORVONERO-PHASE-6.1-CAMPAIGN-FAMILIES-v2.json")],
        ["D-61-02", "Phase 6.1", "Launch geography Novosibirsk + oblast primary-only", "All campaigns", "ACTIVE", rel(PILOTS / "CORVONERO-PHASE-6.1-GEOGRAPHY-ARCHITECTURE-v2.json")],
        ["D-64-01", "Phase 6.4", "LP-01 production content pack approved", "LP-01", "SUPERSEDED by v3", rel(PILOTS / "CORVONERO-PHASE-6.4-LP01-RESULT-v1.json")],
        ["D-66-01", "Phase 6.6", "LP-01 final copy v3 operator approved", "LP-01 public copy", "ACTIVE", rel(PILOTS / "CORVONERO-PHASE-6.6-LP01-FINAL-COPY-APPROVAL-v1.json")],
        ["D-7A-01", "Phase 7A", "LP-01 Tilda staging prep — build authority frozen", "LP-01 implementation", "ACTIVE", rel(PILOTS / "CORVONERO-PHASE-7A-LP01-BUILD-AUTHORITY-MANIFEST-v1.json")],
    ]
    write_sheet(wb.create_sheet("22_DECISION_LOG"), [
        "Decision ID", "Date / phase", "Topic", "Decision", "Affected scope", "Status", "Source",
    ], dec_rows)

    # 23_EXPORT_READINESS
    exp_rows = [
        ["D1 Ads DOCX", "NOT_READY", "ACCEPT registry, campaign map", "Ad copy, negatives, extensions", "Ad authoring wave"],
        ["D2 Landing-page DOCX", "READY — LP-01 exported", "LP-01 final copy v3", "LP-02..LP-06 final copy", "Roman Tilda build LP-01"],
        ["D3 Commander XLSX", "NOT_READY", "Campaign architecture", "Ads, URLs, negatives, import profile", "After ads + published URLs"],
        ["D4 Research XLSX", "CREATED — PARTIAL COVERAGE", "Wordstat, semantics, SERP partial, LP research", "Full SERP, full semantic backlog", "This workbook"],
    ]
    write_sheet(wb.create_sheet("23_EXPORT_READINESS"), [
        "Deliverable", "Current readiness", "Available inputs", "Missing inputs", "Next action",
    ], exp_rows)

    wb.save(path)

    # Reconciliation
    alloc_ids = set(allocation.keys())
    recon = {
        "canonical": counts["canonical"],
        "assessed": counts["assessed"],
        "accept": counts["accept"],
        "reject": counts["reject"],
        "abstain": counts["abstain"],
        "backlog": counts["backlog"],
        "allocation_rows": len(alloc_rows),
        "allocation_unique": len(alloc_ids),
        "campaign_phrase_sum": sum(c.get("allocated_phrase_count", 0) for c in campaigns["campaign_families"]),
        "serp_completed_official": len(official_completed),
        "serp_planned": len(serp_index["queries"]),
    }
    return recon


def verify_docx(path: Path) -> bool:
    try:
        with zipfile.ZipFile(path) as z:
            return "[Content_Types].xml" in z.namelist() and "word/document.xml" in z.namelist()
    except Exception:
        return False


def main():
    if os.environ.get("CORVONERO_OPERATOR_GATE") != "APPROVED":
        raise SystemExit(
            "STOP: CORVONERO_OPERATOR_GATE=APPROVED required. "
            "This script is not safe for casual execution."
        )

    OUT.mkdir(parents=True, exist_ok=True)
    docx_path = OUT / DOCX_NAME
    xlsx_path = OUT / XLSX_NAME

    build_lp01_docx(docx_path)
    recon = build_workbook(xlsx_path)

    docx_ok = verify_docx(docx_path)
    xlsx_ok = xlsx_path.exists() and xlsx_path.stat().st_size > 10000

    outputs = []
    for fp in [docx_path, xlsx_path]:
        outputs.append({
            "path": str(fp),
            "file_type": fp.suffix.lstrip(".").upper(),
            "size_bytes": fp.stat().st_size,
            "sha256": sha256_file(fp),
        })

    validation = {
        "docx_opens": docx_ok,
        "xlsx_exists": xlsx_ok,
        "reconciliation": recon,
        "reconciliation_pass": (
            recon["canonical"] == 2368
            and recon["accept"] == 935
            and recon["reject"] == 368
            and recon["abstain"] == 296
            and recon["backlog"] == 769
            and recon["assessed"] == 1599
            and recon["allocation_rows"] == 935
            and recon["allocation_unique"] == 935
            and recon["campaign_phrase_sum"] == 935
            and recon["serp_completed_official"] == 5
            and recon["serp_planned"] == 10
        ),
    }

    manifest = {
        "manifest_id": "corvonero-export-wave-1-manifest-v1",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "export_date": EXPORT_DATE,
        "source_checkpoint": {
            "commit": CHECKPOINT,
            "tag": "corvonero-pre-export-production-2026-06",
        },
        "outputs": outputs,
        "source_files": [
            rel(PILOTS / "CORVONERO-PHASE-6.6-LP01-FINAL-PRODUCTION-COPY-v3.json"),
            rel(PILOTS / "CORVONERO-PHASE-6.6-LP01-FINAL-FAQ-v3.json"),
            rel(ORCA / "semantic-core/corvonero-canonical-phrase-registry-v1.json"),
            rel(PILOTS / "CORVONERO-RUN-004-PHASE-5.2-FINAL-REVIEWED-REGISTRY-v1.json"),
            rel(PILOTS / "CORVONERO-PHASE-6.1-PHRASE-ALLOCATION-v2.json"),
            rel(MIG / "serp_r1_index.json"),
            rel(PILOTS / "CORVONERO-EXPORT-READINESS-MATRIX-v1.json"),
        ],
        "validation": validation,
        "known_limitations": [
            "Partial semantic coverage 67.5%",
            "SERP 5/10 official boundary",
            "No ads or Commander export",
            "Websites unchanged",
        ],
    }

    manifest_path = OUT / "CORVONERO-EXPORT-WAVE-1-MANIFEST-v1.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    sha_lines = []
    for item in outputs + [{"path": str(manifest_path), "sha256": sha256_file(manifest_path)}]:
        sha_lines.append(f"{item['sha256']}  {Path(item['path']).name}")
    sha_path = OUT / "CORVONERO-EXPORT-WAVE-1-SHA256-v1.txt"
    sha_path.write_text("\n".join(sha_lines) + "\n", encoding="utf-8")

    readme = f"""# CORVONERO Export Wave 1

**Date:** {EXPORT_DATE}  
**Checkpoint:** `{CHECKPOINT}` (`corvonero-pre-export-production-2026-06`)

## Outputs

| File | Purpose |
|------|---------|
| `{DOCX_NAME}` | LP-01 public copy for Roman (Tilda assembly) |
| `{XLSX_NAME}` | Consolidated research and semantic evidence |

## Coverage boundary

- Canonical corpus: **2368**
- Assessed: **1599** (67.5%)
- Unprocessed backlog: **769**
- SERP: **5 / 10** planned queries (official checkpoint boundary)

## Validation

- DOCX verified: **{docx_ok}**
- XLSX verified: **{xlsx_ok}**
- Reconciliation pass: **{validation['reconciliation_pass']}**

## Not included

- D1 Ads DOCX
- D3 Commander XLSX
- Website changes
- Advertising launch
"""
    readme_path = OUT / "CORVONERO-EXPORT-WAVE-1-README-v1.md"
    readme_path.write_text(readme, encoding="utf-8")

    print(json.dumps({"validation": validation, "outputs": [o["path"] for o in outputs]}, ensure_ascii=False, indent=2))
    if not validation["reconciliation_pass"]:
        raise SystemExit("Reconciliation failed")
    if not docx_ok:
        raise SystemExit("DOCX verification failed")


if __name__ == "__main__":
    main()
